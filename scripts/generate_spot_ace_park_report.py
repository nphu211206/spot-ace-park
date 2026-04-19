from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TEMPLATE_PATH = Path(r"C:\Users\Admin\Downloads\Baocao_HUBT_SOCIALdocx.docx")
OUTPUT_DIR = Path(r"C:\Users\Admin\spot-ace-park\docs")
OUTPUT_PATH = OUTPUT_DIR / "Ket_noi_mang_luoi_bai_do_xe_tram_sac_thong_minh_HUBT.docx"

TITLE = "Kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh"
PRODUCT_NAME = "Spot Ace Park"
AUTHOR = "Nguyễn Phú Toàn"
TEAM = [
    "Nguyễn Phú Toàn - TH29.11",
    "Khổng Văn Đản - TH29.19",
    "Nguyễn Quốc Đại - TH29.20",
]
ADVISOR = "ThS. Bùi Thu Giang"


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_cell_border(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str = "Normal",
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    size: float | None = None,
    first_line_cm: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if first_line_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_cm)
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)


def add_body(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        add_paragraph(
            document,
            text,
            style="Normal",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            size=13,
            first_line_cm=1.0,
            space_after=3,
        )


def add_toc_line(document: Document, text: str, style: str) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_box_figure(document: Document, lines: list[str], caption: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    cell = table.cell(0, 0)
    set_cell_border(cell)
    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        set_run_font(run, size=12)
    add_paragraph(
        document,
        caption,
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=12,
        space_before=3,
        space_after=6,
    )


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]], column_widths_cm: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        if column_widths_cm:
            hdr_cells[idx].width = Cm(column_widths_cm[idx])
        paragraph = hdr_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=12, bold=True)
        set_cell_border(hdr_cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if column_widths_cm:
                cells[idx].width = Cm(column_widths_cm[idx])
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_run_font(run, size=12)
            set_cell_border(cells[idx])


def add_page_break(document: Document) -> None:
    document.add_page_break()


INTRO_SECTIONS = {}
CHAPTER_1 = {}
CHAPTER_2 = {}
CHAPTER_3 = {}
CHAPTER_4 = {}
REFERENCES: list[str] = []


def build_cover(document: Document) -> None:
    add_paragraph(document, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=2)
    add_paragraph(document, "TRƯỜNG ĐẠI HỌC KINH DOANH VÀ CÔNG NGHỆ HÀ NỘI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    for _ in range(6):
        add_paragraph(document, "", size=12)
    add_paragraph(document, "BÁO CÁO KẾT QUẢ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    add_paragraph(document, "ĐỀ TÀI NGHIÊN CỨU KHOA HỌC ỨNG DỤNG CẤP TRƯỜNG", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_paragraph(document, "", size=12)
    add_paragraph(document, f"Tên đề tài: {TITLE}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_before=6)
    for _ in range(7):
        add_paragraph(document, "", size=12)
    add_paragraph(document, "Hà Nội - Năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=14)

    add_page_break(document)

    add_paragraph(document, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=2)
    add_paragraph(document, "TRƯỜNG ĐẠI HỌC KINH DOANH VÀ CÔNG NGHỆ HÀ NỘI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    for _ in range(4):
        add_paragraph(document, "", size=12)
    add_paragraph(document, "BÁO CÁO KẾT QUẢ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    add_paragraph(document, "ĐỀ TÀI NGHIÊN CỨU KHOA HỌC ỨNG DỤNG CẤP TRƯỜNG", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_paragraph(document, "", size=12)
    add_paragraph(document, f"Tên đề tài: {TITLE}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_paragraph(document, "", size=12)
    add_paragraph(document, f"Sản phẩm nghiên cứu: Nền tảng {PRODUCT_NAME} tích hợp dự báo chỗ đỗ, điều phối mạng lưới bãi đỗ và trạm sạc thông minh", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=13)
    for _ in range(2):
        add_paragraph(document, "", size=12)
    add_paragraph(document, f"Chủ nhiệm đề tài: {AUTHOR}", size=13, space_after=2)
    add_paragraph(document, "Nhóm thực hiện:", size=13, space_after=0)
    for member in TEAM:
        add_paragraph(document, member, size=13, first_line_cm=0.5, space_after=0)
    add_paragraph(document, "", size=12)
    add_paragraph(document, f"GV hướng dẫn: {ADVISOR}", size=13, space_after=2)
    add_paragraph(document, "Đơn vị thực hiện: Khoa Công nghệ thông tin", size=13, space_after=2)
    add_paragraph(document, "Thời gian hoàn thành báo cáo: tháng 4 năm 2026", size=13)
    for _ in range(4):
        add_paragraph(document, "", size=12)
    add_paragraph(document, "Hà Nội - Năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=14)


def build_front_matter(document: Document) -> None:
    add_page_break(document)
    add_paragraph(document, "MỤC LỤC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    toc_lines = [
        ("DANH MỤC HÌNH\t4", "toc 1"),
        ("LỜI MỞ ĐẦU\t5", "toc 1"),
        ("1. Tính cấp thiết của đề tài\t5", "toc 2"),
        ("2. Tổng quan tình hình nghiên cứu\t6", "toc 2"),
        ("2.1 Ngoài nước\t6", "toc 3"),
        ("2.2 Trong nước\t7", "toc 3"),
        ("3. Mục tiêu nghiên cứu\t8", "toc 2"),
        ("4. Đối tượng phạm vi nghiên cứu\t8", "toc 2"),
        ("5. Nội dung nghiên cứu\t9", "toc 2"),
        ("6. Phương pháp nghiên cứu và tiếp cận vấn đề\t10", "toc 2"),
        ("7. Ý nghĩa và tính mới về khoa học thực tiễn\t11", "toc 2"),
        ("7.1 Mục tiêu\t11", "toc 3"),
        ("7.2 Xác định các yếu tố ảnh hưởng\t11", "toc 3"),
        ("7.3 Phát triển mô hình điều phối dự báo\t12", "toc 3"),
        ("7.4 Giá trị mang lại\t12", "toc 3"),
        ("7.5 Tính đổi mới\t13", "toc 3"),
        ("8. Kết cấu đề tài\t14", "toc 2"),
        ("CHƯƠNG 1: CƠ SỞ LÝ LUẬN\t15", "toc 1"),
        ("1.1. Khái niệm và vai trò của mạng lưới bãi đỗ xe - trạm sạc thông minh\t15", "toc 2"),
        ("1.1.1. Khái niệm hệ thống\t15", "toc 3"),
        ("1.1.2. Vai trò của hệ thống đối với đô thị và người dùng\t16", "toc 3"),
        ("1.2. Cơ sở lý thuyết về dự báo nhu cầu, chấp nhận công nghệ và song sinh số\t17", "toc 2"),
        ("1.2.1. Lý thuyết dự báo nhu cầu và điều phối tài nguyên\t17", "toc 3"),
        ("1.2.2. Lý thuyết hành vi chấp nhận công nghệ\t18", "toc 3"),
        ("1.2.3. Lý thuyết song sinh số trong quản trị vận hành\t18", "toc 3"),
        ("1.3. Cơ sở pháp lý và xu hướng phát triển hạ tầng\t19", "toc 2"),
        ("1.3.1. Xu hướng phát triển hạ tầng sạc công cộng\t19", "toc 3"),
        ("1.3.2. Định hướng pháp lý đối với bãi đỗ xe và trạm sạc\t20", "toc 3"),
        ("1.3.3. Giao thức tích hợp trạm sạc và dữ liệu vận hành\t20", "toc 3"),
        ("1.3.4. Ý nghĩa của việc kết nối bãi đỗ với trạm sạc\t21", "toc 3"),
        ("CHƯƠNG 2: THỰC TRẠNG, ĐÁNH GIÁ THỰC TRẠNG\t22", "toc 1"),
        ("2.1. Thực trạng nhu cầu tìm chỗ đỗ xe và trạm sạc tại đô thị\t22", "toc 2"),
        ("2.1.1. Các khó khăn phổ biến của người dùng\t22", "toc 3"),
        ("2.1.2. Khó khăn của người dùng xe điện\t23", "toc 3"),
        ("2.1.3. Mức độ tiếp cận các công cụ hiện có\t24", "toc 3"),
        ("2.2. Thực trạng công tác vận hành của đơn vị khai thác bãi xe\t25", "toc 2"),
        ("2.3. Đánh giá thực trạng và khoảng trống cần khắc phục\t26", "toc 2"),
        ("CHƯƠNG 3: GIẢI PHÁP, TRIỂN KHAI THỰC HIỆN\t27", "toc 1"),
        ("3.1. Mô hình tổng quan hệ thống\t27", "toc 2"),
        ("3.2. Sơ đồ Usecase\t28", "toc 2"),
        ("3.2.1. Sơ đồ usecase của người dùng và quản lý bãi xe\t28", "toc 3"),
        ("3.2.2. Sơ đồ usecase của Admin\t29", "toc 3"),
        ("3.3. Các tính năng chính của ứng dụng\t30", "toc 2"),
        ("3.3.1. Đối với người dùng\t30", "toc 3"),
        ("3.3.2. Đối với quản lý bãi xe\t31", "toc 3"),
        ("3.3.3. Đối với Admin\t32", "toc 3"),
        ("3.4. Kiến trúc kỹ thuật và công nghệ sử dụng\t33", "toc 2"),
        ("3.5. Thiết kế cơ sở dữ liệu (CSDL)\t35", "toc 2"),
        ("3.6. Thiết kế giao diện người dùng (UI/UX)\t36", "toc 2"),
        ("3.7. Các luồng hoạt động chính của ứng dụng\t40", "toc 2"),
        ("3.8. Các biện pháp bảo mật hệ thống\t43", "toc 2"),
        ("CHƯƠNG 4: KẾT LUẬN, KIẾN NGHỊ\t45", "toc 1"),
        ("4.1. Những điểm chính mà đề tài đã đạt được\t45", "toc 2"),
        ("4.2. Những điểm mà đề tài chưa hoàn thiện được và hướng giải quyết\t46", "toc 2"),
        ("4.2.1. Những điểm chưa hoàn thiện\t46", "toc 3"),
        ("4.2.2. Hướng giải quyết\t47", "toc 3"),
        ("4.3. Hướng phát triển của đề tài\t48", "toc 2"),
        ("4.4. Kết luận chung\t49", "toc 2"),
        ("TÀI LIỆU THAM KHẢO\t50", "toc 1"),
    ]
    for text, style in toc_lines:
        add_toc_line(document, text, style)

    add_page_break(document)
    add_heading(document, "DANH MỤC HÌNH", 1)
    figure_lines = [
        "Hình 3.1 Mô hình tổng quan hệ thống Spot Ace Park\t27",
        "Hình 3.2 Sơ đồ usecase của người dùng và quản lý bãi xe\t28",
        "Hình 3.3 Sơ đồ usecase của Admin\t29",
        "Hình 3.4 Luồng đặt chỗ theo ETA và cơ chế cam kết chỗ đỗ\t40",
        "Hình 3.5 Luồng park-and-charge cho xe điện\t41",
        "Hình 3.6 Mô hình digital twin tái cấu hình bãi đỗ\t42",
        "Hình 3.7 Giao diện đăng nhập\t36",
        "Hình 3.8 Giao diện bản đồ mạng lưới bãi đỗ\t37",
        "Hình 3.9 Giao diện đặt chỗ dự báo\t38",
        "Hình 3.10 Giao diện quản lý phiên đỗ và sạc\t39",
        "Hình 3.11 Giao diện dashboard quản lý bãi xe\t39",
        "Hình 3.12 Giao diện digital twin cho vận hành\t40",
    ]
    for line in figure_lines:
        add_toc_line(document, line, "toc 2")


INTRO_SECTIONS = {
    "1. Tính cấp thiết của đề tài": [
        "Sự gia tăng nhanh của phương tiện cá nhân tại các đô thị lớn đang làm cho bài toán tìm chỗ đỗ xe trở thành một vấn đề giao thông đô thị có tác động trực tiếp đến thời gian di chuyển, chi phí xã hội, mức tiêu hao nhiên liệu và trải nghiệm sống của người dân. Trong nhiều trường hợp, người dùng không thiếu chỗ đỗ xe theo nghĩa tuyệt đối, mà thiếu thông tin tin cậy để biết bãi đỗ nào còn chỗ, ở thời điểm nào, cách điểm đến bao xa và có phù hợp với kế hoạch di chuyển hay không.",
        "Bên cạnh áp lực đỗ xe truyền thống, quá trình chuyển dịch sang phương tiện điện đặt ra thêm một lớp nhu cầu mới. Người dùng xe điện không chỉ cần một chỗ đỗ, mà cần một chỗ đỗ có kết nối với hạ tầng sạc phù hợp về công suất, loại đầu nối, tình trạng hoạt động và thời điểm sử dụng. Khi mạng lưới bãi đỗ và trạm sạc chưa được kết nối trong cùng một nền tảng, người dùng thường phải mở nhiều ứng dụng riêng lẻ, dẫn đến thông tin rời rạc, quyết định chậm và rủi ro đến nơi nhưng không còn tài nguyên phù hợp.",
        "Ở góc độ đơn vị vận hành, các bãi đỗ hiện nay vẫn chủ yếu khai thác theo mô hình tĩnh: số lượng ô đỗ, khu vực ưu tiên, mức giá và luồng phương tiện ít thay đổi theo ngữ cảnh. Trong khi đó, nhu cầu thực tế lại biến động mạnh theo khung giờ, thời tiết, sự kiện, khu vực và xu hướng sử dụng xe điện. Nếu không có một hệ thống dự báo và điều phối thông minh, tài nguyên đỗ xe và tài nguyên điện sẽ khó được khai thác tối ưu.",
        "Do đó, đề tài \"Kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh\" mang tính cấp thiết cả về hạ tầng đô thị, trải nghiệm người dùng lẫn hiệu quả vận hành. Điểm trọng tâm của đề tài không chỉ nằm ở việc hiển thị thông tin bãi đỗ hoặc trạm sạc, mà ở khả năng dự báo, giữ chỗ theo thời gian đến dự kiến, điều phối thay thế khi bãi đầy và tái cấu hình tài nguyên đỗ xe - sạc theo thời gian thực trên cùng một nền tảng thống nhất.",
    ],
    "2. Tổng quan tình hình nghiên cứu": [
        "Các nghiên cứu và sản phẩm thực tiễn trên thế giới cho thấy lĩnh vực smart parking đã phát triển theo nhiều nhánh khác nhau như nhận diện ô trống bằng IoT, đặt chỗ trước, định tuyến tới bãi đỗ, thanh toán điện tử và phân tích công suất khai thác. Song song với đó, lĩnh vực quản lý trạm sạc xe điện cũng phát triển mạnh với các hệ thống tìm trạm, quản lý phiên sạc, cân bằng tải và kết nối qua giao thức chuẩn. Tuy nhiên, trong nhiều trường hợp hai bài toán này vẫn tồn tại như hai hệ sinh thái tách rời.",
    ],
    "2.1 Ngoài nước": [
        "Ở thị trường quốc tế, nhiều nền tảng đã cho phép đặt chỗ bãi đỗ hoặc tìm trạm sạc theo bản đồ. Một số nghiên cứu gần đây còn xem bãi đỗ có tích hợp trạm sạc như một tài sản năng lượng, nơi việc phân bổ công suất và hoạch định hạ tầng có thể ảnh hưởng trực tiếp đến hiệu quả đầu tư. Các xu hướng nổi bật gồm dự báo nhu cầu theo thời gian thực, phân loại người dùng theo hành vi đỗ xe, cân bằng tải giữa các trụ sạc và tích hợp dữ liệu vận hành với dashboard điều hành trung tâm.",
        "Tuy vậy, qua khảo sát tài liệu công khai, không phải mọi hệ thống đều giải quyết trọn vẹn bài toán từ góc nhìn hành trình người dùng. Người dùng cần biết không chỉ chỗ nào đang trống ở hiện tại, mà chỗ nào có khả năng còn trống khi họ thực sự đến nơi; nếu tình huống thay đổi thì hệ thống sẽ phản ứng như thế nào; và nếu kết hợp cả đỗ xe lẫn sạc điện thì việc giữ chỗ, bắt đầu phiên sạc và tính phí tổng hợp sẽ được xử lý ra sao.",
    ],
    "2.2 Trong nước": [
        "Tại Việt Nam, thị trường đã xuất hiện một số nhóm giải pháp riêng lẻ như ứng dụng cung cấp thông tin bãi đỗ theo thời gian thực, nền tảng bản đồ trạm sạc, ứng dụng quản lý phiên sạc hoặc các hệ thống quản lý bãi xe tại từng đơn vị. Điều đó cho thấy nhu cầu số hóa là có thật và đang tăng nhanh. Tuy nhiên, các sản phẩm công khai hiện nay phần lớn tập trung vào một lát cắt của bài toán: hoặc là đỗ xe, hoặc là sạc xe điện, hoặc là quản trị hạ tầng sạc.",
        "Khoảng trống đáng chú ý là một nền tảng có thể vừa kết nối mạng lưới bãi đỗ xe và trạm sạc, vừa vận hành theo logic dự báo và điều phối. Nghĩa là hệ thống phải biết tài nguyên nào nên được giữ cho khách đang tới, khu vực nào nên chuyển đổi sang ô EV trong khung giờ cao điểm, và khi dự báo sai thì cơ chế thay thế, hoàn ưu đãi hoặc chuyển hướng cần được thực hiện ra sao. Đây chính là điểm mở để đề tài đề xuất một mô hình mới và có chiều sâu hơn so với các giải pháp tra cứu thông thường.",
    ],
    "3. Mục tiêu nghiên cứu": [
        "Mục tiêu tổng quát của đề tài là xây dựng mô hình và giải pháp phần mềm cho một nền tảng kết nối mạng lưới bãi đỗ xe và trạm sạc thông minh theo hướng dự báo - điều phối - bảo đảm dịch vụ, lấy sản phẩm mẫu Spot Ace Park làm môi trường nghiên cứu và mô phỏng triển khai.",
        "Các mục tiêu cụ thể bao gồm: xây dựng bộ chức năng cho người dùng nhằm tìm kiếm, đặt chỗ và theo dõi tài nguyên đỗ xe - sạc theo thời gian thực; xây dựng dashboard cho quản lý bãi xe và quản trị viên để điều hành, theo dõi công suất và ra quyết định; xây dựng logic dự báo nhu cầu và cơ chế cam kết giữ chỗ theo thời gian đến dự kiến; đề xuất mô hình digital twin để mô phỏng tái cấu hình bãi đỗ và đánh giá hiệu quả vận hành trước khi áp dụng vào thực tế.",
    ],
    "4. Đối tượng phạm vi nghiên cứu": [
        "Đối tượng nghiên cứu của đề tài là hệ thống phần mềm quản lý và điều phối mạng lưới bãi đỗ xe - trạm sạc thông minh trong bối cảnh đô thị, bao gồm ba nhóm tác nhân chính: người dùng phương tiện, đơn vị vận hành bãi xe và quản trị viên hệ thống.",
        "Phạm vi nghiên cứu tập trung vào phần mềm và dữ liệu vận hành, không phụ thuộc bắt buộc vào việc lắp đặt phần cứng thật ở giai đoạn đầu. Hệ thống được thiết kế để có thể mô phỏng dữ liệu camera, barrier, cảm biến ra/vào và trạng thái trạm sạc, từ đó cho phép kiểm chứng các luồng nghiệp vụ và thuật toán điều phối trước khi tích hợp phần cứng thật.",
    ],
    "5. Nội dung nghiên cứu": [
        "Đề tài tập trung vào năm nhóm nội dung chính. Thứ nhất, khảo sát hiện trạng và khoảng trống của các hệ thống đỗ xe và sạc điện. Thứ hai, xây dựng mô hình nghiệp vụ cho nền tảng kết nối bãi đỗ và trạm sạc. Thứ ba, thiết kế kiến trúc kỹ thuật, cơ sở dữ liệu và giao diện người dùng cho sản phẩm Spot Ace Park. Thứ tư, đề xuất thuật toán dự báo nhu cầu và cơ chế giữ chỗ theo ETA kết hợp điều phối thay thế. Thứ năm, xây dựng mô hình digital twin để hỗ trợ quản lý tái cấu hình tài nguyên đỗ xe - sạc theo thời gian thực.",
    ],
    "6. Phương pháp nghiên cứu và tiếp cận vấn đề": [
        "Đề tài sử dụng phương pháp nghiên cứu kết hợp giữa khảo sát tài liệu, phân tích nghiệp vụ, thiết kế hệ thống thông tin, xây dựng nguyên mẫu phần mềm và mô phỏng dữ liệu vận hành. Cách tiếp cận này phù hợp với một đề tài ứng dụng vì vừa có thể bảo vệ phần lý luận, vừa có khả năng hiện thực hóa thành sản phẩm demo hoạt động được.",
        "Bên cạnh phân tích định tính về nhu cầu người dùng và bài toán vận hành, đề tài còn tiếp cận từ góc nhìn hệ thống: xem bãi đỗ và trạm sạc như một mạng lưới tài nguyên có trạng thái biến đổi liên tục theo thời gian. Từ đó, phần mềm phải có khả năng quan sát, dự báo, đưa ra quyết định và phản ứng với tình huống thay đổi chứ không chỉ dừng ở việc hiển thị dữ liệu tĩnh.",
    ],
    "7. Ý nghĩa và tính mới về khoa học thực tiễn": [
        "Về mặt thực tiễn, đề tài hướng tới một mô hình phần mềm có thể giải quyết cùng lúc ba nhu cầu: giảm thời gian tìm chỗ đỗ, kết nối người dùng xe điện với trạm sạc phù hợp, và giúp đơn vị vận hành nâng cao hiệu quả khai thác hạ tầng. Về mặt học thuật, đề tài chuyển trọng tâm từ một ứng dụng tra cứu sang một hệ điều phối vận hành dựa trên dự báo và ra quyết định.",
    ],
    "7.1 Mục tiêu": [
        "Tính mới đầu tiên của đề tài nằm ở việc xác định lại mục tiêu hệ thống: không chỉ cho biết tài nguyên nào đang sẵn có, mà còn dự báo tài nguyên nào sẽ còn phù hợp trong tương lai rất gần đối với từng người dùng cụ thể. Điều này làm thay đổi cách thiết kế chức năng, dữ liệu và thuật toán của hệ thống.",
    ],
    "7.2 Xác định các yếu tố ảnh hưởng": [
        "Đề tài không xem nhu cầu đỗ xe như một biến tĩnh mà xác định nó chịu tác động đồng thời của nhiều yếu tố: khung giờ, loại khu vực, mật độ xe vào ra, thời gian lưu đỗ trung bình, điều kiện thời tiết, sự kiện lân cận, tỷ lệ xe điện, khả năng cấp điện của bãi và hành vi người dùng. Việc nhận diện đầy đủ các yếu tố này là tiền đề để xây dựng mô hình điều phối khả thi hơn.",
    ],
    "7.3 Phát triển mô hình điều phối dự báo": [
        "Tính mới nổi bật nhất của đề tài là mô hình điều phối dự báo gồm bốn lớp: dự báo công suất theo ETA, bộ máy giữ chỗ có điều kiện, cơ chế cam kết dịch vụ và bộ máy thay thế khi tài nguyên bị biến động. Mô hình này cho phép hệ thống phản ứng linh hoạt trước sai số dự báo thay vì đẩy toàn bộ rủi ro sang phía người dùng.",
    ],
    "7.4 Giá trị mang lại": [
        "Giá trị mang lại của đề tài không chỉ ở việc tạo ra một giao diện hiện đại, mà ở khả năng làm nền tảng cho ra quyết định vận hành. Người dùng được hỗ trợ ở cấp độ hành trình; quản lý bãi xe được hỗ trợ ở cấp độ tài nguyên; còn quản trị viên được hỗ trợ ở cấp độ quy hoạch và tối ưu mạng lưới.",
    ],
    "7.5 Tính đổi mới": [
        "Qua khảo sát các giải pháp công khai hiện có, nhóm nghiên cứu chưa ghi nhận phổ biến một hệ thống tại Việt Nam mô tả đồng thời bốn yếu tố: dự báo chỗ trống theo thời gian đến, cơ chế cam kết giữ chỗ và điều phối thay thế, mô hình số tái cấu hình bãi đỗ trong cùng dashboard vận hành, và tích hợp park-and-charge như một dịch vụ thống nhất. Đây là cơ sở để khẳng định tính đổi mới của đề tài ở mức mô hình giải pháp và sản phẩm ứng dụng.",
    ],
    "8. Kết cấu đề tài": [
        "Ngoài phần mở đầu và tài liệu tham khảo, báo cáo được bố cục thành bốn chương. Chương 1 trình bày cơ sở lý luận và các khung tham chiếu liên quan đến smart parking, EV charging, dự báo nhu cầu và song sinh số. Chương 2 phân tích thực trạng và khoảng trống của các giải pháp hiện hành. Chương 3 trình bày mô hình hệ thống, tính năng, kiến trúc kỹ thuật, thiết kế dữ liệu, giao diện và các luồng vận hành chính. Chương 4 tổng kết kết quả, nêu hạn chế và định hướng phát triển tiếp theo.",
    ],
}


CHAPTER_1 = {
    "1.1. Khái niệm và vai trò của mạng lưới bãi đỗ xe - trạm sạc thông minh": [
        "Mạng lưới bãi đỗ xe - trạm sạc thông minh là hệ thống kết nối nhiều điểm đỗ xe, khu vực đỗ, trụ sạc và thiết bị giám sát vào cùng một nền tảng dữ liệu, cho phép theo dõi trạng thái tài nguyên theo thời gian thực, cung cấp dịch vụ đặt chỗ, điều phối và thanh toán trên cùng một hạ tầng phần mềm.",
        "Khác với mô hình bãi xe truyền thống vận hành cục bộ, mô hình mạng lưới cho phép xem mỗi bãi đỗ là một nút trong một hệ sinh thái rộng hơn. Khi đó, quyết định điều phối không chỉ nhằm tối ưu cho từng bãi riêng lẻ mà còn tối ưu cho cả khu vực phục vụ. Người dùng có thể được gợi ý bãi thay thế nếu bãi đích quá tải; đơn vị quản lý có thể san tải nhu cầu giữa các bãi và phân phối tài nguyên EV theo thực trạng điện năng từng điểm.",
    ],
    "1.1.1. Khái niệm hệ thống": [
        "Trong phạm vi đề tài này, hệ thống được hiểu là tổ hợp gồm ba lớp: lớp tương tác người dùng, lớp dịch vụ điều phối nghiệp vụ và lớp dữ liệu - tích hợp thiết bị. Lớp tương tác bao gồm ứng dụng cho người dùng, dashboard quản lý và dashboard quản trị. Lớp dịch vụ chịu trách nhiệm xử lý tìm kiếm, đặt chỗ, dự báo, điều phối và thanh toán. Lớp dữ liệu - tích hợp thiết bị quản lý trạng thái ô đỗ, log vào ra, trạng thái trạm sạc, phiên sạc, định danh phương tiện và dữ liệu sự kiện từ hiện trường hoặc bộ mô phỏng.",
    ],
    "1.1.2. Vai trò của hệ thống đối với đô thị và người dùng": [
        "Đối với người dùng, hệ thống làm giảm chi phí tìm kiếm thông tin và rủi ro ra quyết định trong môi trường giao thông nhiều biến động. Đối với đơn vị khai thác, hệ thống là công cụ ra quyết định giúp tăng công suất sử dụng tài sản, quản trị doanh thu và giảm thất thoát vận hành. Đối với đô thị, hệ thống góp phần hạn chế lưu lượng xe chạy vòng quanh tìm chỗ đỗ, hỗ trợ hình thành hạ tầng giao thông thông minh và thúc đẩy quá trình chuyển đổi sang phương tiện điện.",
    ],
    "1.2. Cơ sở lý thuyết về dự báo nhu cầu, chấp nhận công nghệ và song sinh số": [
        "Đề tài tham chiếu ba nhóm cơ sở lý thuyết chính. Thứ nhất là lý thuyết dự báo nhu cầu và điều phối tài nguyên, nền tảng cho việc ước lượng khả năng còn chỗ trong tương lai gần và phân bổ tài nguyên cho người dùng. Thứ hai là lý thuyết hành vi chấp nhận công nghệ, giúp giải thích vì sao người dùng lựa chọn một nền tảng số khi nền tảng đó mang lại giá trị về tiện lợi, độ tin cậy và minh bạch. Thứ ba là lý thuyết song sinh số, cho phép mô phỏng trạng thái hệ thống vật lý bằng mô hình số để hỗ trợ thử nghiệm quyết định trước khi áp dụng vào thực tế.",
    ],
    "1.2.1. Lý thuyết dự báo nhu cầu và điều phối tài nguyên": [
        "Bài toán cốt lõi của nền tảng không phải là kiểm đếm chỗ trống tại thời điểm hiện tại, mà là ước lượng khả năng chỗ trống còn phù hợp tại một thời điểm trong tương lai gần. Cách nhìn này đưa đề tài từ mô hình giám sát sang mô hình dự báo. Khi có dự báo, hệ thống mới có thể gán trước tài nguyên cho khách đang di chuyển đến và chủ động điều phối khi có biến động.",
    ],
    "1.2.2. Lý thuyết hành vi chấp nhận công nghệ": [
        "Một ứng dụng chỉ thực sự được chấp nhận khi người dùng thấy rằng nó hữu ích, dễ sử dụng và đáng tin cậy. Trong bối cảnh đỗ xe và sạc điện, yếu tố đáng tin cậy đặc biệt quan trọng. Nếu hệ thống nhiều lần báo còn chỗ nhưng đến nơi lại hết chỗ, người dùng sẽ nhanh chóng mất niềm tin. Vì vậy, thiết kế của đề tài đặt trọng tâm vào cơ chế cam kết dịch vụ và phản ứng khi dự báo sai để bảo vệ niềm tin đó.",
    ],
    "1.2.3. Lý thuyết song sinh số trong quản trị vận hành": [
        "Song sinh số là một bản sao số của hệ thống vật lý, cho phép theo dõi trạng thái hiện tại, mô phỏng tình huống và đánh giá tác động của một quyết định. Áp dụng vào bãi đỗ xe - trạm sạc, digital twin giúp quản lý thử các kịch bản như tăng số ô EV, thay đổi vùng ưu tiên, hạn chế công suất sạc hoặc điều chỉnh giá theo giờ trước khi ban hành trong vận hành thật.",
    ],
    "1.3. Cơ sở pháp lý và xu hướng phát triển hạ tầng": [
        "Ngoài các luận điểm công nghệ, đề tài còn dựa trên xu hướng phát triển hạ tầng và chính sách giao thông thông minh. Hạ tầng sạc công cộng đang mở rộng nhanh ở nhiều quốc gia, trong đó khu vực Đông Nam Á cũng ghi nhận tốc độ tăng trưởng đáng kể. Ở Việt Nam, Luật Đường bộ 2024 bắt đầu định hình khung pháp lý mới cho giao thông thông minh, bến xe, trạm dừng nghỉ và hệ thống sạc điện cho phương tiện sử dụng năng lượng điện.",
    ],
    "1.3.1. Xu hướng phát triển hạ tầng sạc công cộng": [
        "Theo IEA trong báo cáo Global EV Outlook 2025, hạ tầng sạc công cộng tại nhiều thị trường đang tiếp tục tăng trưởng mạnh, đặc biệt ở nhóm quốc gia mới nổi. Riêng Indonesia, Thailand, Malaysia và Việt Nam đã có hơn 24.000 bộ sạc công cộng, cao gấp 9 lần so với năm 2022. Điều này cho thấy nhu cầu tổ chức, quản trị và kết nối dữ liệu giữa trạm sạc với hệ sinh thái giao thông đang trở nên cấp thiết.",
    ],
    "1.3.2. Định hướng pháp lý đối với bãi đỗ xe và trạm sạc": [
        "Luật Đường bộ số 35/2024/QH15 có hiệu lực từ ngày 01/01/2025 nhấn mạnh định hướng giao thông thông minh, ứng dụng công nghệ và tích hợp dữ liệu trong quản lý, khai thác kết cấu hạ tầng đường bộ. Luật cũng nêu yêu cầu trạm dừng nghỉ, bến xe phải có hệ thống sạc điện cho phương tiện giao thông cơ giới sử dụng năng lượng điện và ưu tiên bố trí hệ thống sạc tại bãi đỗ xe. Đây là cơ sở pháp lý quan trọng để khẳng định sự phù hợp của hướng nghiên cứu.",
    ],
    "1.3.3. Giao thức tích hợp trạm sạc và dữ liệu vận hành": [
        "Về kỹ thuật tích hợp, Open Charge Alliance xác định OCPP 2.0.1 như một hướng chuẩn quan trọng cho kết nối giữa hệ thống quản lý trạm sạc và thiết bị sạc. OCPP 2.0.1 mở rộng mạnh về quản lý thiết bị, an toàn, xử lý giao dịch và smart charging, rất phù hợp để làm tài liệu tham chiếu khi đề tài đề xuất khả năng kết nối với trạm sạc thật trong giai đoạn sau.",
    ],
    "1.3.4. Ý nghĩa của việc kết nối bãi đỗ với trạm sạc": [
        "Nếu bãi đỗ và trạm sạc được xem là hai hạ tầng riêng biệt, người dùng vẫn phải tự giải bài toán nối hai hành trình: đỗ xe ở đâu và sạc xe ở đâu. Khi hai hạ tầng này được kết nối ngay từ lớp dữ liệu và dịch vụ, hệ thống có thể nhìn toàn bộ chuỗi nhu cầu của người dùng, từ đó đưa ra quyết định đồng bộ hơn, ví dụ giữ ô đỗ EV phù hợp, tính trước thời gian sạc cần thiết hoặc tối ưu thời điểm cấp điện.",
    ],
}


CHAPTER_2 = {
    "2.1. Thực trạng nhu cầu tìm chỗ đỗ xe và trạm sạc tại đô thị": [
        "Trong môi trường đô thị, nhu cầu đỗ xe chịu ảnh hưởng rất lớn của vị trí, thời gian và loại mục đích di chuyển. Cùng một bãi xe có thể thừa chỗ vào buổi sáng nhưng quá tải vào khung giờ tối; cùng một khu vực có thể thuận lợi cho xe xăng nhưng chưa phù hợp với xe điện vì thiếu chỗ đỗ EV hoặc thiếu công suất sạc. Điều này khiến hành vi lựa chọn bãi đỗ không chỉ là bài toán khoảng cách, mà là bài toán đa tiêu chí.",
    ],
    "2.1.1. Các khó khăn phổ biến của người dùng": [
        "Người dùng thường gặp ba khó khăn chính. Một là thông tin phân tán, phải tự so sánh giữa nhiều kênh. Hai là thông tin không cập nhật kịp thời, dẫn đến độ tin cậy thấp. Ba là thiếu cơ chế bảo đảm: ngay cả khi ứng dụng báo còn chỗ, người dùng cũng không biết chỗ đó có còn phù hợp khi họ đến nơi hay không. Điều này làm cho nhiều nền tảng chỉ có giá trị tham khảo chứ chưa trở thành công cụ ra quyết định thực sự.",
        "Một trở ngại nữa là trải nghiệm vận hành sau khi đến bãi xe còn rời rạc. Người dùng có thể phải qua nhiều bước thủ công như lấy vé, xác nhận biển số, thanh toán riêng, hoặc không có kênh theo dõi thời gian đỗ. Khi nhu cầu tăng cao, sự thiếu liền mạch này trực tiếp làm giảm chất lượng dịch vụ.",
    ],
    "2.1.2. Khó khăn của người dùng xe điện": [
        "Đối với người dùng xe điện, khó khăn còn cao hơn vì họ phải cân nhắc đồng thời vị trí đỗ xe, loại đầu nối, công suất sạc, thời lượng lưu đỗ và khả năng có hàng chờ. Nhiều trường hợp bãi xe có trạm sạc nhưng không còn ô EV; hoặc ô EV còn trống nhưng trụ sạc đang lỗi, đang bảo trì hoặc công suất không đáp ứng nhu cầu. Nếu không có nền tảng tích hợp, người dùng rất khó đưa ra phương án tối ưu cho chuyến đi.",
    ],
    "2.1.3. Mức độ tiếp cận các công cụ hiện có": [
        "Khảo sát công khai cho thấy thị trường đã có ứng dụng đặt chỗ bãi đỗ, có ứng dụng tìm trạm sạc và có nền tảng quản lý trạm sạc. Tuy nhiên, mức độ liên thông dữ liệu và chiều sâu điều phối còn hạn chế. Đa số công cụ mạnh ở khâu tra cứu trạng thái hiện tại, trong khi bài toán người dùng quan tâm lại là trạng thái tại thời điểm họ đến nơi. Khoảng chênh này là lý do khiến nhiều người vẫn dựa vào thói quen hoặc kinh nghiệm cá nhân hơn là hoàn toàn tin vào phần mềm.",
    ],
    "2.2. Thực trạng công tác vận hành của đơn vị khai thác bãi xe": [
        "Ở phía đơn vị vận hành, khó khăn thường nằm ở ba tầng: dữ liệu, điều hành và quy hoạch. Về dữ liệu, thông tin xe vào ra, công suất ô đỗ, trạng thái trụ sạc và doanh thu nhiều khi nằm ở các hệ thống tách biệt. Về điều hành, quản lý khó phản ứng nhanh khi một khu vực đột ngột quá tải, khi xe điện tăng cao hoặc khi trụ sạc gặp sự cố. Về quy hoạch, đơn vị khai thác thiếu công cụ mô phỏng để biết nên mở rộng ô EV ở đâu, thời điểm nào hoặc bao nhiêu là hợp lý.",
        "Nếu chỉ dựa vào số liệu thống kê cuối ngày hoặc báo cáo thủ công, đơn vị vận hành sẽ khó tối ưu tài nguyên theo thời gian thực. Điều này dẫn đến lãng phí ở cả hai đầu: chỗ đỗ có thể bị khóa cho mục đích không phù hợp, trong khi khách hàng có nhu cầu thực lại không được phục vụ đúng lúc.",
    ],
    "2.3. Đánh giá thực trạng và khoảng trống cần khắc phục": [
        "Từ các phân tích trên có thể thấy khoảng trống lớn nhất hiện nay không phải là thiếu ứng dụng, mà là thiếu một lớp điều phối thông minh nằm giữa người dùng và hạ tầng. Lớp điều phối này phải có khả năng dự báo nhu cầu theo ETA, quyết định ai được giữ chỗ nào, điều phối thay thế ra sao khi tình huống thay đổi và gợi ý tái cấu hình tài nguyên cho đơn vị vận hành.",
        "Chính vì vậy, đề tài tập trung vào việc xây dựng Spot Ace Park như một nền tảng kết nối và điều phối. Tính ứng dụng của đề tài đến từ việc sản phẩm có thể chạy với dữ liệu mô phỏng ngay trong giai đoạn nghiên cứu, đồng thời vẫn giữ khả năng mở rộng để tích hợp phần cứng và trạm sạc thật trong giai đoạn triển khai sau này.",
    ],
}


CHAPTER_3 = {
    "3.1. Mô hình tổng quan hệ thống": [
        "Mô hình tổng quan của Spot Ace Park được xây dựng theo hướng nền tảng dịch vụ nhiều vai trò. Ở lớp đầu vào là người dùng, quản lý bãi xe và quản trị viên hệ thống. Ở lớp xử lý trung tâm là các dịch vụ tìm kiếm, đặt chỗ, dự báo ETA, định giá, điều phối thay thế, quản lý phiên sạc và digital twin. Ở lớp dữ liệu là các thực thể bãi đỗ, ô đỗ, trụ sạc, phương tiện, đặt chỗ, log vào ra và cấu hình vận hành. Cách tổ chức này giúp hệ thống vừa đáp ứng trải nghiệm người dùng, vừa hỗ trợ vận hành và mở rộng kỹ thuật.",
        "Điểm trung tâm của mô hình là Reservation & Guarantee Engine. Đây là bộ máy quyết định xem tài nguyên nào nên được giữ cho người dùng đang di chuyển tới, trong bao lâu, với mức độ ưu tiên nào, và nếu tài nguyên đó biến động thì phương án thay thế nào là hợp lý nhất. Bằng cách này, hệ thống chuyển từ cơ chế hiển thị trạng thái sang cơ chế quản trị cam kết dịch vụ.",
    ],
    "3.2. Sơ đồ Usecase": [
        "Usecase của hệ thống được chia làm hai nhóm lớn. Nhóm thứ nhất bao gồm người dùng và quản lý bãi xe, tập trung vào hành trình sử dụng và vận hành hiện trường. Nhóm thứ hai là quản trị viên, tập trung vào giám sát, cấu hình và quản lý toàn mạng lưới.",
    ],
    "3.2.1. Sơ đồ usecase của người dùng và quản lý bãi xe": [
        "Người dùng có thể đăng ký tài khoản, khai báo phương tiện, tìm bãi đỗ phù hợp, lọc bãi có trạm sạc, xem xác suất còn chỗ theo ETA, đặt chỗ, check-in bằng QR hoặc biển số, theo dõi phiên đỗ xe và phiên sạc, thanh toán và xem lịch sử giao dịch. Quản lý bãi xe có thể theo dõi công suất, duyệt cấu hình khu vực, giám sát log vào ra, theo dõi hàng chờ, theo dõi phiên sạc, xem cảnh báo và nhận đề xuất tái cấu hình bãi.",
    ],
    "3.2.2. Sơ đồ usecase của Admin": [
        "Admin có vai trò quản trị toàn hệ thống: quản lý người dùng và phân quyền, quản lý danh mục bãi đỗ và trạm sạc, cấu hình thuật toán, xem báo cáo tổng hợp, giám sát hiệu năng mạng lưới, cấu hình giá dịch vụ, quản lý nhật ký hệ thống và các tham số dùng cho engine dự báo.",
    ],
    "3.3. Các tính năng chính của ứng dụng": [
        "Các tính năng của ứng dụng được thiết kế theo ba lớp vai trò, trong đó lớp người dùng phục vụ hành trình thực tế, lớp quản lý bãi xe phục vụ vận hành tác nghiệp và lớp admin phục vụ điều hành chiến lược và cấu hình hệ thống.",
    ],
    "3.3.1. Đối với người dùng": [
        "Đối với người dùng, các tính năng cốt lõi gồm: tìm kiếm bãi đỗ theo vị trí và điểm đến; xem công suất còn lại theo thời gian thực; xem dự báo xác suất còn chỗ tại thời điểm đến dự kiến; đặt chỗ theo ETA; xem đề xuất bãi thay thế; chọn ô EV hoặc trụ sạc phù hợp; check-in bằng QR hoặc biển số; theo dõi đồng hồ đỗ xe và phiên sạc; thanh toán hợp nhất; đánh giá chất lượng bãi đỗ.",
        "Điểm khác biệt lớn nhất là trải nghiệm đặt chỗ không dừng ở thao tác chọn ô đỗ, mà đi kèm một lớp bảo đảm dịch vụ. Hệ thống có thể hiển thị độ tin cậy, thời gian giữ chỗ, điều kiện hủy giữ chỗ và phương án chuyển hướng nếu tài nguyên ban đầu thay đổi quá ngưỡng cho phép.",
    ],
    "3.3.2. Đối với quản lý bãi xe": [
        "Đối với quản lý bãi xe, các tính năng chính bao gồm: dashboard theo dõi xe đang đỗ, lượng xe ra vào, doanh thu, công suất theo khu; theo dõi trạng thái ô EV, trụ sạc và phiên sạc; xem log cảm biến hoặc camera mô phỏng; quản lý hàng chờ; nhận cảnh báo quá tải; áp dụng gợi ý tái cấu hình số ô đỗ thường và ô EV; xem báo cáo theo ca, theo ngày và theo sự kiện.",
    ],
    "3.3.3. Đối với Admin": [
        "Đối với admin, hệ thống cung cấp: quản lý mạng lưới bãi đỗ và trạm sạc; quản lý danh mục người dùng và quyền hạn; cấu hình engine dự báo; cấu hình quy tắc giá và ngưỡng cam kết; theo dõi hiệu năng các bãi trên bản đồ; xem nhật ký thay đổi cấu hình; xuất báo cáo tổng hợp và mô phỏng mở rộng hạ tầng ở cấp mạng lưới.",
    ],
    "3.4. Kiến trúc kỹ thuật và công nghệ sử dụng": [
        "Kiến trúc kỹ thuật của hệ thống được chia thành frontend, backend dịch vụ, tầng dữ liệu và tầng tích hợp. Frontend sử dụng React, TypeScript và Vite để xây dựng giao diện người dùng. Backend sử dụng Node.js và Express cho REST API, xử lý nghiệp vụ và tích hợp dữ liệu thời gian thực. SQL Server lưu trữ dữ liệu nghiệp vụ chính. Lớp mô phỏng hoặc tích hợp thiết bị tiếp nhận dữ liệu từ camera, cảm biến, barrier và trạm sạc. Lớp realtime có thể dùng WebSocket hoặc dịch vụ đồng bộ sự kiện để cập nhật trạng thái hệ thống ngay trên dashboard.",
        "Sản phẩm mẫu hiện tại của đề tài đã có nền giao diện quản lý, đặt chỗ, thanh toán, OCR biển số và dashboard mô phỏng. Trên nền đó, bản final mở rộng thêm mô hình reservation theo ETA, cơ chế guarantee, quản lý hàng chờ EV và digital twin. Điều này giúp đề tài có tính kế thừa và khả năng triển khai kỹ thuật thực sự, thay vì chỉ dừng ở mức ý tưởng khái niệm.",
    ],
    "3.5. Thiết kế cơ sở dữ liệu (CSDL)": [
        "Thiết kế cơ sở dữ liệu được mở rộng từ các bảng hiện có của hệ thống thành một mô hình phù hợp hơn với mạng lưới bãi đỗ và trạm sạc. Nhóm bảng cốt lõi gồm AppUsers, Vehicles, ParkingLots, ParkingZones, ParkingSlots, Bookings, ChargingStations, ChargingConnectors, ChargingReservations, ChargingSessions, IoTLogs, Payments, Notifications, SystemConfig và AuditLogs.",
        "Trong đó, các bảng liên quan đến trạm sạc không chỉ lưu trạng thái thiết bị mà còn lưu khả năng công suất, loại đầu nối, khu vực lắp đặt, trạng thái vận hành và các chỉ số phục vụ điều phối. Bảng Bookings và ChargingReservations được liên kết để tạo thành dịch vụ park-and-charge, giúp người dùng có thể được phục vụ trên một hành trình thống nhất thay vì hai giao dịch rời rạc.",
    ],
    "3.6. Thiết kế giao diện người dùng (UI/UX)": [
        "UI/UX của Spot Ace Park được thiết kế theo nguyên tắc trực quan, có thứ bậc thông tin rõ ràng và hỗ trợ ra quyết định nhanh. Người dùng cần thấy ngay các yếu tố quan trọng nhất gồm: khoảng cách, giá, xác suất còn chỗ, mức độ phù hợp EV và thời gian đi bộ đến điểm đến. Quản lý bãi xe cần thấy công suất, cảnh báo, luồng xe và đề xuất tái cấu hình. Admin cần có bản đồ tổng quan và báo cáo đa chiều.",
    ],
    "3.6.1. Logo của ứng dụng": [
        "Logo đề xuất của Spot Ace Park kết hợp biểu tượng vị trí, ô đỗ và tia năng lượng để thể hiện hai lớp dịch vụ: đỗ xe và sạc điện. Màu chủ đạo thiên về xanh dương kết hợp xanh lá nhằm gợi tính công nghệ, tin cậy và phát triển bền vững.",
    ],
    "3.6.2. Giao diện đăng nhập": [
        "Giao diện đăng nhập tập trung vào nhận diện thương hiệu và trải nghiệm vào hệ thống nhanh, hỗ trợ phân quyền theo vai trò người dùng, quản lý bãi xe và admin. Phần xác thực có thể mở rộng sang OTP hoặc SSO trong giai đoạn triển khai thật.",
    ],
    "3.6.3. Giao diện bản đồ mạng lưới bãi đỗ": [
        "Giao diện bản đồ mạng lưới hiển thị danh sách bãi đỗ và trạm sạc theo khu vực, cho phép lọc theo giá, khoảng cách, xác suất còn chỗ, loại phương tiện và nhu cầu sạc. Đây là giao diện giúp người dùng ra quyết định trước chuyến đi.",
    ],
    "3.6.4. Giao diện đặt chỗ dự báo": [
        "Giao diện đặt chỗ dự báo là màn hình đặc trưng của đề tài. Tại đây, người dùng không chỉ nhìn thấy tài nguyên hiện tại mà còn nhìn thấy khả năng tài nguyên còn phù hợp tại thời điểm đến. Màn hình hiển thị ETA, độ tin cậy, thời gian giữ chỗ và điều kiện cam kết dịch vụ.",
    ],
    "3.6.5. Giao diện park-and-charge cho xe điện": [
        "Màn hình park-and-charge cho phép người dùng chọn chỗ đỗ có sạc, xem loại đầu nối, công suất, thời gian sạc dự kiến, mức phí sạc và thời gian giữ chỗ EV. Trong cùng một giao diện, hệ thống cho phép kết hợp đặt chỗ đỗ và đặt chỗ sạc thành một luồng duy nhất.",
    ],
    "3.6.6. Giao diện quản lý phiên đỗ và sạc": [
        "Màn hình quản lý phiên hiển thị đồng hồ đếm thời gian, trạng thái check-in, thời gian còn lại, mức phí tạm tính, trạng thái phiên sạc, năng lượng đã tiêu thụ và các thông báo như 80% pin, sắp hết thời gian giữ chỗ hoặc cảnh báo sắp phát sinh idle fee.",
    ],
    "3.6.7. Giao diện dashboard quản lý bãi xe": [
        "Dashboard quản lý bãi xe ưu tiên thông tin thời gian thực: số xe đang đỗ, công suất theo khu, tỷ lệ ô EV đang sử dụng, số xe đang chờ, trạng thái trụ sạc và bản đồ nhiệt theo giờ. Giao diện này đóng vai trò trung tâm điều hành của bãi.",
    ],
    "3.6.8. Giao diện digital twin cho vận hành": [
        "Giao diện digital twin cho phép quản lý xem mô hình số của bãi xe, thay đổi giả lập số ô thường, ô EV, vùng ưu tiên và giới hạn công suất để đánh giá tác động tới doanh thu, công suất và mức độ phục vụ. Đây là phần thể hiện rõ nhất tính mới của đề tài ở góc độ vận hành thông minh.",
    ],
    "3.6.9. Giao diện báo cáo quản trị mạng lưới": [
        "Giao diện báo cáo quản trị mạng lưới cung cấp cái nhìn đa bãi, cho phép so sánh công suất, tỷ lệ đặt chỗ thành công, tỷ lệ fallback, mức sử dụng trụ sạc, doanh thu theo bãi và các chỉ số SLA phục vụ cho quyết định mở rộng hoặc điều chỉnh chính sách.",
    ],
    "3.7. Các luồng hoạt động chính của ứng dụng": [
        "Để làm rõ khả năng ứng dụng của hệ thống, đề tài mô tả năm luồng nghiệp vụ tiêu biểu phản ánh đầy đủ tư tưởng thiết kế của Spot Ace Park final.",
    ],
    "3.7.1. Luồng đặt chỗ theo ETA và cơ chế cam kết chỗ đỗ": [
        "Người dùng nhập điểm đến và thời gian dự kiến tới. Hệ thống tính ETA, đối chiếu dữ liệu công suất lịch sử và trạng thái hiện tại để dự báo xác suất còn chỗ ở từng bãi. Sau khi người dùng chọn một bãi phù hợp, Reservation & Guarantee Engine ghi nhận yêu cầu, giữ chỗ trong một khoảng thời gian có điều kiện và liên tục kiểm tra lại trạng thái bãi. Nếu xác suất rơi xuống dưới ngưỡng an toàn, hệ thống chủ động gửi phương án thay thế trước khi người dùng tới nơi.",
    ],
    "3.7.2. Luồng check-in và điều phối vào bãi": [
        "Khi người dùng đến bãi, hệ thống check-in bằng QR hoặc nhận diện biển số. Nếu người dùng đến đúng trong thời gian giữ chỗ, barrier được mở và phiên đỗ bắt đầu. Nếu người dùng đến muộn, hệ thống đánh giá lại điều kiện giữ chỗ dựa trên công suất thực tế. Luồng này cho phép cân bằng giữa quyền lợi của khách đặt trước và hiệu quả sử dụng tài nguyên chung.",
    ],
    "3.7.3. Luồng park-and-charge cho xe điện": [
        "Đối với xe điện, hệ thống ngoài việc giữ một ô đỗ EV còn có thể gắn kèm một tài nguyên sạc cụ thể. Người dùng chọn loại đầu nối, công suất mong muốn hoặc mức pin mục tiêu. Hệ thống tạo một bản ghi booking gắn với một charging reservation. Sau khi check-in, phiên đỗ và phiên sạc được mở đồng thời; người dùng theo dõi toàn bộ hành trình trên một màn hình duy nhất.",
    ],
    "3.7.4. Luồng điều phối thay thế khi bãi đầy": [
        "Nếu bãi đích sắp đầy do biến động đột xuất, hệ thống sẽ tìm bãi thay thế trong cùng khu vực dựa trên các tiêu chí ưu tiên: khoảng cách, thời gian đi bộ, trạng thái còn chỗ, mức độ tương thích EV và mức độ ảnh hưởng tới ETA. Khi người dùng chấp nhận phương án thay thế, booking ban đầu được cập nhật và hệ thống chuyển dữ liệu điều hướng sang bãi mới. Đây là cơ chế biến cam kết dịch vụ thành một hành động vận hành cụ thể.",
    ],
    "3.7.5. Luồng tái cấu hình bãi đỗ bằng digital twin": [
        "Ở vai trò quản lý, digital twin nhận dữ liệu hiện trạng của bãi và cho phép giả lập các kịch bản. Ví dụ, nếu tăng số ô EV từ 6 lên 10 trong khung giờ 17h-21h thì tỷ lệ phục vụ xe điện tăng bao nhiêu, ô thường giảm bao nhiêu, doanh thu thay đổi như thế nào, công suất sạc có vượt ngưỡng không. Khi một kịch bản được phê duyệt, hệ thống ghi nhận cấu hình mới và dùng làm cơ sở điều phối ở chu kỳ tiếp theo.",
    ],
    "3.8. Các biện pháp bảo mật hệ thống": [
        "Hệ thống áp dụng phân quyền theo vai trò, ghi nhật ký thay đổi cấu hình, kiểm soát truy cập API, xác thực người dùng, mã hóa dữ liệu nhạy cảm và tách biệt quyền quản trị mạng lưới khỏi quyền vận hành từng bãi. Các sự kiện quan trọng như thay đổi giá, đổi cấu hình ô EV, mở quyền admin hoặc can thiệp phiên sạc đều phải được ghi vết để phục vụ audit.",
        "Trong giai đoạn triển khai thực tế, hệ thống có thể mở rộng thêm chữ ký giao dịch, chuẩn hóa webhook từ thiết bị, xác thực đa yếu tố cho quản trị viên và cơ chế cảnh báo khi dữ liệu hiện trường có dấu hiệu bất thường. Cấu trúc này cho phép sản phẩm phát triển từ bản nghiên cứu sang môi trường vận hành thực mà không cần thay đổi lại toàn bộ kiến trúc bảo mật.",
    ],
}


CHAPTER_4 = {
    "4.1. Những điểm chính mà đề tài đã đạt được": [
        "Đề tài đã hệ thống hóa được bài toán kết nối bãi đỗ xe với trạm sạc dưới góc nhìn điều phối thay vì chỉ tra cứu. Trên cơ sở đó, đề tài đề xuất được một mô hình sản phẩm cụ thể với tên gọi Spot Ace Park, có cấu trúc dữ liệu, kiến trúc kỹ thuật, giao diện người dùng và các luồng nghiệp vụ đủ rõ để chuyển sang pha triển khai phần mềm.",
        "Một kết quả quan trọng khác là xác định được điểm khác biệt cốt lõi của sản phẩm so với nhiều giải pháp hiện hành: dự báo công suất theo ETA, cơ chế cam kết giữ chỗ, điều phối thay thế khi bãi đầy và digital twin phục vụ tái cấu hình tài nguyên. Đây là nền tảng để đề tài giữ được tính mới mà vẫn bảo đảm khả năng thực thi kỹ thuật.",
    ],
    "4.2. Những điểm mà đề tài chưa hoàn thiện được và hướng giải quyết": [
        "Do giới hạn về thời gian và phạm vi nghiên cứu, đề tài chưa tích hợp phần cứng thật như barrier, camera ANPR chuyên dụng hoặc trạm sạc vật lý. Một số chức năng nâng cao như dự báo bằng mô hình máy học thời gian thực, cân bằng tải điện năng chi tiết hoặc tích hợp dữ liệu giao thông đô thị cũng mới dừng ở mức thiết kế và mô phỏng.",
    ],
    "4.2.1. Những điểm chưa hoàn thiện": [
        "Các thiếu hụt chính của giai đoạn hiện tại gồm: chưa có dữ liệu vận hành diện rộng đủ lớn để huấn luyện mô hình dự báo nâng cao; chưa hoàn tất đầy đủ luồng tích hợp park-and-charge trên backend; chưa có bộ chỉ số đánh giá SLA qua nhiều bãi thật; và chưa triển khai thử nghiệm với đơn vị khai thác ngoài phạm vi mô phỏng.",
    ],
    "4.2.2. Hướng giải quyết": [
        "Để khắc phục, nhóm nghiên cứu định hướng triển khai theo từng lớp. Lớp thứ nhất là hoàn thiện phần mềm lõi trên dữ liệu mô phỏng và dữ liệu nội bộ. Lớp thứ hai là tích hợp dần thiết bị và trạm sạc ở quy mô pilot. Lớp thứ ba là tối ưu mô hình dự báo bằng dữ liệu thực tế và mở rộng digital twin sang công cụ hỗ trợ quy hoạch đầu tư.",
    ],
    "4.3. Hướng phát triển của đề tài": [
        "Trong giai đoạn tiếp theo, đề tài có thể mở rộng theo ba hướng. Một là kết nối dữ liệu giao thông đô thị và sự kiện khu vực để tăng độ chính xác dự báo ETA và công suất bãi. Hai là triển khai smart charging và điều phối công suất theo giới hạn điện năng thực. Ba là phát triển module quy hoạch đầu tư, cho phép chủ bãi mô phỏng chi phí - doanh thu - mức phục vụ khi tăng số ô EV hoặc số trụ sạc.",
    ],
    "4.4. Kết luận chung": [
        "Đề tài \"Kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh\" đã xác lập được một hướng tiếp cận có tính ứng dụng và tính mới rõ ràng: biến bãi đỗ và trạm sạc từ hạ tầng rời rạc thành một mạng lưới có khả năng quan sát, dự báo, điều phối và tự tối ưu theo bối cảnh sử dụng. Sản phẩm Spot Ace Park được đề xuất như một nền tảng phần mềm có thể làm cầu nối giữa người dùng, đơn vị vận hành và hạ tầng thông minh trong đô thị hiện đại.",
        "Với cách tiếp cận software-first, đề tài không bị phụ thuộc tuyệt đối vào việc có phần cứng thật ngay từ đầu nhưng vẫn bảo đảm khả năng chuyển tiếp sang triển khai thực tế. Đây là cơ sở để nhóm tiếp tục phát triển phần mềm trong giai đoạn sau, đồng thời bảo vệ được định hướng khoa học ứng dụng của công trình nghiên cứu.",
    ],
}


REFERENCES = [
    "[1] International Energy Agency (IEA), Global EV Outlook 2025, mục Electric vehicle charging, 2025.",
    "[2] Open Charge Alliance, OCPP 2.0.1 Protocol Overview, truy cập 2026.",
    "[3] Open Charge Alliance, What is new in OCPP 2.0.1?, whitepaper, 2023.",
    "[4] Luật số 35/2024/QH15 của Quốc hội: Luật Đường bộ, ban hành ngày 27/06/2024, có hiệu lực từ 01/01/2025.",
    "[5] Bộ Tư pháp, Những nội dung cơ bản của Luật Đường bộ số 35/2024/QH15, 2024.",
    "[6] Smart car parks with EV charging for academic campus, ScienceDirect, 2024.",
    "[7] V-Green, thông tin công khai về hạ tầng sạc và nền tảng tích hợp đối tác, truy cập 2026.",
    "[8] My Parking, thông tin mô tả ứng dụng cung cấp trạng thái bãi đỗ theo thời gian thực và đặt chỗ, truy cập 2026.",
    "[9] Việt EV App, thông tin mô tả tính năng bản đồ trạm sạc, trạng thái thời gian thực và trip planner, truy cập 2026.",
]


def build_chapter_1(document: Document) -> None:
    add_heading(document, "CHƯƠNG 1: CƠ SỞ LÝ LUẬN", 1)
    for heading, paragraphs in CHAPTER_1.items():
        level = 3 if heading.startswith(("1.1.1", "1.1.2", "1.2.1", "1.2.2", "1.2.3", "1.3.1", "1.3.2", "1.3.3", "1.3.4")) else 2
        add_heading(document, heading, level)
        add_body(document, paragraphs)


def build_chapter_2(document: Document) -> None:
    add_heading(document, "CHƯƠNG 2: THỰC TRẠNG, ĐÁNH GIÁ THỰC TRẠNG", 1)
    for heading, paragraphs in CHAPTER_2.items():
        level = 3 if heading.startswith(("2.1.1", "2.1.2", "2.1.3")) else 2
        add_heading(document, heading, level)
        add_body(document, paragraphs)

    add_paragraph(document, "Bảng 2.1. So sánh nhanh một số nhóm giải pháp công khai trên thị trường", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, space_before=4, space_after=4)
    add_simple_table(
        document,
        ["Nhóm giải pháp", "Điểm mạnh chính", "Khoảng trống còn lại"],
        [
            ["Ứng dụng bãi đỗ", "Realtime, đặt chỗ, thanh toán cho một số điểm đỗ", "Chưa tập trung sâu vào dự báo ETA và cam kết giữ chỗ"],
            ["Ứng dụng bản đồ trạm sạc", "Tìm trạm, xem đầu nối, trạng thái, điều hướng", "Ít gắn với tài nguyên đỗ xe và vận hành bãi"],
            ["Nền tảng quản lý sạc", "Giám sát phiên sạc, thiết bị, giao dịch", "Thiếu góc nhìn hành trình đỗ xe - sạc thống nhất cho người dùng"],
            ["Spot Ace Park đề xuất", "Kết nối bãi đỗ, trạm sạc, ETA, guarantee, fallback, digital twin", "Cần tiếp tục tích hợp thiết bị thật và dữ liệu pilot ở giai đoạn sau"],
        ],
        [4.0, 5.5, 7.0],
    )


def build_chapter_3(document: Document) -> None:
    add_heading(document, "CHƯƠNG 3: GIẢI PHÁP, TRIỂN KHAI THỰC HIỆN", 1)
    order = [
        "3.1. Mô hình tổng quan hệ thống",
        "3.2. Sơ đồ Usecase",
        "3.2.1. Sơ đồ usecase của người dùng và quản lý bãi xe",
        "3.2.2. Sơ đồ usecase của Admin",
        "3.3. Các tính năng chính của ứng dụng",
        "3.3.1. Đối với người dùng",
        "3.3.2. Đối với quản lý bãi xe",
        "3.3.3. Đối với Admin",
        "3.4. Kiến trúc kỹ thuật và công nghệ sử dụng",
        "3.5. Thiết kế cơ sở dữ liệu (CSDL)",
        "3.6. Thiết kế giao diện người dùng (UI/UX)",
        "3.6.1. Logo của ứng dụng",
        "3.6.2. Giao diện đăng nhập",
        "3.6.3. Giao diện bản đồ mạng lưới bãi đỗ",
        "3.6.4. Giao diện đặt chỗ dự báo",
        "3.6.5. Giao diện park-and-charge cho xe điện",
        "3.6.6. Giao diện quản lý phiên đỗ và sạc",
        "3.6.7. Giao diện dashboard quản lý bãi xe",
        "3.6.8. Giao diện digital twin cho vận hành",
        "3.6.9. Giao diện báo cáo quản trị mạng lưới",
        "3.7. Các luồng hoạt động chính của ứng dụng",
        "3.7.1. Luồng đặt chỗ theo ETA và cơ chế cam kết chỗ đỗ",
        "3.7.2. Luồng check-in và điều phối vào bãi",
        "3.7.3. Luồng park-and-charge cho xe điện",
        "3.7.4. Luồng điều phối thay thế khi bãi đầy",
        "3.7.5. Luồng tái cấu hình bãi đỗ bằng digital twin",
        "3.8. Các biện pháp bảo mật hệ thống",
    ]
    for heading in order:
        level = 3 if heading.startswith(("3.2.1", "3.2.2", "3.3.1", "3.3.2", "3.3.3", "3.6.1", "3.6.2", "3.6.3", "3.6.4", "3.6.5", "3.6.6", "3.6.7", "3.6.8", "3.6.9", "3.7.1", "3.7.2", "3.7.3", "3.7.4", "3.7.5")) else 2
        add_heading(document, heading, level)
        add_body(document, CHAPTER_3[heading])

        if heading == "3.1. Mô hình tổng quan hệ thống":
            add_box_figure(document, ["Người dùng / Quản lý bãi / Admin", "                |", "Frontend Web App & Dashboard", "                |", "Reservation & Guarantee Engine", "                |", "Forecasting - Pricing - Fallback - Digital Twin", "                |", "Parking Network - EV Charging Network - IoT Events - Payments"], "Hình 3.1. Mô hình tổng quan hệ thống Spot Ace Park")
        elif heading == "3.2.1. Sơ đồ usecase của người dùng và quản lý bãi xe":
            add_box_figure(document, ["Người dùng:", "Đăng nhập | Tìm bãi | Xem ETA | Đặt chỗ | Chọn EV slot | Check-in | Theo dõi phiên | Thanh toán", "", "Quản lý bãi xe:", "Theo dõi công suất | Quản lý khu đỗ | Giám sát trụ sạc | Xem log vào/ra | Nhận cảnh báo | Tái cấu hình bãi"], "Hình 3.2. Sơ đồ usecase của người dùng và quản lý bãi xe")
        elif heading == "3.2.2. Sơ đồ usecase của Admin":
            add_box_figure(document, ["Admin", "Quản lý người dùng | Quản lý bãi đỗ | Quản lý trạm sạc", "Cấu hình engine | Theo dõi mạng lưới | Audit log | Báo cáo tổng hợp"], "Hình 3.3. Sơ đồ usecase của Admin")
        elif heading == "3.4. Kiến trúc kỹ thuật và công nghệ sử dụng":
            add_paragraph(document, "Bảng 3.1. Công nghệ sử dụng trong đề tài", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, space_before=4, space_after=4)
            add_simple_table(document, ["Nhóm", "Công nghệ / Định hướng"], [["Frontend", "React, TypeScript, Vite, Tailwind CSS, shadcn/ui, Framer Motion"], ["Backend", "Node.js, Express, REST API, mô phỏng realtime bằng WebSocket hoặc event stream"], ["CSDL", "SQL Server cho dữ liệu nghiệp vụ chính; có thể mở rộng cache hoặc queue khi triển khai lớn"], ["Nhận diện", "OCR/ANPR cho biển số xe, QR check-in/check-out"], ["EV Charging", "Mô hình dữ liệu charger/session, tham chiếu khả năng tích hợp OCPP 2.0.1"], ["Vận hành", "Dashboard manager/admin, digital twin, báo cáo và audit log"]], [4.0, 12.5])
        elif heading == "3.5. Thiết kế cơ sở dữ liệu (CSDL)":
            add_paragraph(document, "Bảng 3.2. Nhóm bảng dữ liệu chính của hệ thống", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, space_before=4, space_after=4)
            add_simple_table(document, ["Nhóm bảng", "Vai trò"], [["AppUsers, Vehicles", "Quản lý người dùng, phương tiện, loại xe, khả năng EV, biển số"], ["ParkingLots, ParkingZones, ParkingSlots", "Mô hình hóa mạng lưới bãi đỗ, khu vực và tài nguyên ô đỗ"], ["Bookings", "Lưu giao dịch đặt chỗ, thời gian giữ chỗ, trạng thái check-in/check-out"], ["ChargingStations, ChargingConnectors", "Lưu trạm sạc, đầu nối, công suất, trạng thái thiết bị"], ["ChargingReservations, ChargingSessions", "Quản lý giữ chỗ sạc và phiên sạc thực tế"], ["IoTLogs, Notifications, AuditLogs", "Ghi log sự kiện, gửi cảnh báo và lưu vết thao tác hệ thống"], ["Payments, SystemConfig", "Quản lý thanh toán, cấu hình giá, ngưỡng engine và tham số vận hành"]], [5.0, 11.5])
        elif heading == "3.6.2. Giao diện đăng nhập":
            add_box_figure(document, ["LOGO SPOT ACE PARK", "", "[ Ô nhập số điện thoại ]", "[ Ô nhập mật khẩu ]", "[ Nút đăng nhập ]    [ Điều hướng đăng ký ]", "", "Phân quyền: User / Manager / Admin"], "Hình 3.7. Giao diện đăng nhập")
        elif heading == "3.6.3. Giao diện bản đồ mạng lưới bãi đỗ":
            add_box_figure(document, ["BẢN ĐỒ KHU VỰC", "Bộ lọc: Khoảng cách | Giá | EV | Xác suất còn chỗ", "Danh sách bãi đỗ: Tên bãi | ETA | Giá | Độ tin cậy | Số chỗ EV"], "Hình 3.8. Giao diện bản đồ mạng lưới bãi đỗ")
        elif heading == "3.6.4. Giao diện đặt chỗ dự báo":
            add_box_figure(document, ["BÃI A - 650m - ETA 12 phút", "Xác suất còn chỗ khi đến: 92%", "Giá tạm tính: 30.000đ/giờ", "Giữ chỗ tối đa: 15 phút từ thời điểm ETA", "[ Chọn bãi này ]    [ Xem bãi thay thế ]"], "Hình 3.9. Giao diện đặt chỗ dự báo")
        elif heading == "3.6.6. Giao diện quản lý phiên đỗ và sạc":
            add_box_figure(document, ["PHIÊN ĐỖ XE / SẠC ĐIỆN", "Biển số: 29A-123.45 | Ô EV: E-05 | Trụ sạc: CCS2-02", "Thời gian còn lại: 01:25:00", "Điện năng đã nạp: 14.6 kWh | Chi phí tạm tính: 98.000đ", "Thông báo: đạt 80% pin | sắp phát sinh idle fee"], "Hình 3.10. Giao diện quản lý phiên đỗ và sạc")
        elif heading == "3.6.7. Giao diện dashboard quản lý bãi xe":
            add_box_figure(document, ["KPI: Xe đang đỗ | Công suất | Doanh thu | Xe EV đang phục vụ", "Biểu đồ công suất theo giờ", "Trạng thái trụ sạc theo khu vực", "Gợi ý tái cấu hình: +4 ô EV tại khu B từ 17h đến 21h"], "Hình 3.11. Giao diện dashboard quản lý bãi xe")
        elif heading == "3.6.8. Giao diện digital twin cho vận hành":
            add_box_figure(document, ["DIGITAL TWIN", "Kịch bản A: 60 ô thường | 10 ô EV | 4 trụ sạc", "Kịch bản B: 54 ô thường | 16 ô EV | 6 trụ sạc", "So sánh: Công suất phục vụ | Doanh thu | Tỷ lệ fallback | Tải điện"], "Hình 3.12. Giao diện digital twin cho vận hành")
        elif heading == "3.7.1. Luồng đặt chỗ theo ETA và cơ chế cam kết chỗ đỗ":
            add_box_figure(document, ["Nhập điểm đến -> Tính ETA -> Dự báo còn chỗ -> Chọn bãi", "-> Reservation Engine giữ chỗ có điều kiện", "-> Theo dõi biến động -> Xác nhận khi check-in"], "Hình 3.4. Luồng đặt chỗ theo ETA và cơ chế cam kết chỗ đỗ")
        elif heading == "3.7.3. Luồng park-and-charge cho xe điện":
            add_box_figure(document, ["Chọn bãi EV -> Chọn đầu nối/công suất -> Giữ ô EV + trụ sạc", "-> Check-in bằng QR/biển số -> Mở phiên đỗ + sạc", "-> Theo dõi điện năng -> Thanh toán hợp nhất"], "Hình 3.5. Luồng park-and-charge cho xe điện")
        elif heading == "3.7.5. Luồng tái cấu hình bãi đỗ bằng digital twin":
            add_box_figure(document, ["Dữ liệu hiện trạng -> Mô hình số của bãi", "-> Giả lập thay đổi số ô thường / ô EV / trụ sạc", "-> So sánh KPI -> Phê duyệt cấu hình mới"], "Hình 3.6. Mô hình digital twin tái cấu hình bãi đỗ")


def build_chapter_4(document: Document) -> None:
    add_heading(document, "CHƯƠNG 4: KẾT LUẬN, KIẾN NGHỊ", 1)
    for heading, paragraphs in CHAPTER_4.items():
        level = 3 if heading.startswith(("4.2.1", "4.2.2")) else 2
        add_heading(document, heading, level)
        add_body(document, paragraphs)


def build_references(document: Document) -> None:
    add_heading(document, "TÀI LIỆU THAM KHẢO", 1)
    for ref in REFERENCES:
        add_paragraph(document, ref, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, space_after=2)


def build_document() -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(str(TEMPLATE_PATH))
    clear_document(document)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(13)

    build_cover(document)
    build_front_matter(document)

    add_page_break(document)
    add_heading(document, "LỜI MỞ ĐẦU", 1)
    for heading, paragraphs in INTRO_SECTIONS.items():
        level = 3 if heading.startswith(("2.1", "2.2", "7.1", "7.2", "7.3", "7.4", "7.5")) else 2
        add_heading(document, heading, level)
        add_body(document, paragraphs)

    add_page_break(document)
    build_chapter_1(document)
    add_page_break(document)
    build_chapter_2(document)
    add_page_break(document)
    build_chapter_3(document)
    add_page_break(document)
    build_chapter_4(document)
    add_page_break(document)
    build_references(document)

    document.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
