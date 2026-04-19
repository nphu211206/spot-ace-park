from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_PATH = Path(r"C:\Users\Admin\Downloads\Thuyet_Minh_Nguyen_Phu_Toan_TH29.11-3 (1).docx")
OUTPUT_DIR = Path(r"C:\Users\Admin\spot-ace-park\docs")
OUTPUT_PATH = OUTPUT_DIR / "Thuyet_Minh_Ket_noi_mang_luoi_bai_do_xe_tram_sac_thong_minh_giu_nguyen_form_v3.docx"


GOAL_TEXT = (
    "Đề tài tập trung xây dựng nền tảng Spot Ace Park theo định hướng kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh "
    "trong cùng một hệ thống phần mềm thống nhất.\n"
    "Mục tiêu chính của đề tài là chuyển bài toán từ mức tra cứu thông tin sang mức dự báo, điều phối và bảo đảm dịch vụ.\n"
    "- Đối với người dùng: tìm bãi đỗ theo điểm đến, xem xác suất còn chỗ theo thời gian đến dự kiến (ETA), đặt chỗ có điều kiện, "
    "nhận gợi ý bãi thay thế khi bãi đích biến động.\n"
    "- Đối với người dùng xe điện: chọn ô đỗ có sạc, đầu nối phù hợp, theo dõi phiên park-and-charge và thanh toán gộp.\n"
    "- Đối với quản lý bãi xe: theo dõi công suất, trạng thái khu đỗ/ô EV/trụ sạc, nhận cảnh báo quá tải và gợi ý tái cấu hình.\n"
    "- Đối với quản trị viên: quản lý mạng lưới bãi đỗ - trạm sạc, cấu hình engine dự báo và theo dõi hiệu năng toàn hệ thống.\n"
    "- Về mặt công nghệ: phát triển Reservation & Guarantee Engine, fallback routing, dashboard quản trị và digital twin cho vận hành.\n"
    "- Về định hướng triển khai: áp dụng mô hình software-first, có thể mô phỏng dữ liệu camera, barrier và charger trước khi tích hợp thiết bị thật."
)

OVERVIEW_TEXT = (
    "6.1. Tổng quan tình hình nghiên cứu:\n"
    "Trên thế giới, smart parking và quản lý trạm sạc xe điện đều đã phát triển mạnh với các hướng như IoT nhận diện ô trống, "
    "đặt chỗ trước, bản đồ trạm sạc, quản lý phiên sạc và smart charging. Tuy nhiên, nhiều hệ thống vẫn xử lý bài toán đỗ xe và "
    "bài toán sạc điện như hai mảng tách rời. Người dùng thường phải tự ghép thông tin từ nhiều nguồn khác nhau để quyết định đi đâu, đỗ ở đâu và sạc thế nào.\n"
    "Tại Việt Nam, các giải pháp công khai hiện nay chủ yếu đi theo từng lát cắt riêng: ứng dụng bãi đỗ, bản đồ trạm sạc, nền tảng quản lý charger hoặc hệ thống nội bộ của từng bãi xe. "
    "Khoảng trống còn lại là một lớp điều phối đủ thông minh để nhìn toàn bộ hành trình sử dụng tài nguyên.\n\n"
    "6.2. Tính cấp thiết của đề tài:\n"
    "Áp lực giao thông đô thị, thời gian tìm chỗ đỗ xe kéo dài, xu hướng gia tăng xe điện và yêu cầu tích hợp hạ tầng sạc đang làm cho mô hình bãi đỗ tĩnh trở nên kém hiệu quả. "
    "Người dùng không chỉ cần biết bãi nào đang trống, mà cần biết bãi nào có khả năng còn phù hợp khi họ đến nơi. "
    "Đơn vị vận hành cũng cần công cụ dự báo và điều phối thay vì chỉ theo dõi số liệu hiện tại.\n\n"
    "6.3. Ý nghĩa và tính mới về khoa học và thực tiễn:\n"
    "Điểm mới của đề tài nằm ở mô hình kết nối bãi đỗ - trạm sạc theo hướng ETA-based reservation, cơ chế cam kết giữ chỗ, "
    "điều phối bãi thay thế khi bãi đầy và digital twin phục vụ tái cấu hình tài nguyên. "
    "Về thực tiễn, đề tài giúp giảm thời gian tìm chỗ, tăng khả năng phục vụ xe điện và nâng cao hiệu quả khai thác bãi xe."
)

RESEARCH_CONTENT_TEXT = (
    "Nội dung nghiên cứu\n"
    "7.1. Phân tích và thiết kế hệ thống:\n"
    "- Phân tích yêu cầu chức năng và phi chức năng cho ba nhóm vai trò: người dùng, quản lý bãi xe, admin.\n"
    "- Thiết kế kiến trúc hệ thống theo mô hình web app + dashboard + backend điều phối.\n"
    "- Thiết kế mô hình dữ liệu cho bãi đỗ, ô đỗ, trụ sạc, reservation, charging session, IoT logs và audit logs.\n"
    "7.2. Xây dựng module cho người dùng:\n"
    "- Tìm kiếm bãi đỗ theo điểm đến, khoảng cách, giá và hỗ trợ EV.\n"
    "- Xem dự báo xác suất còn chỗ theo ETA và đặt chỗ có điều kiện.\n"
    "- Theo dõi phiên đỗ xe, phiên sạc và thanh toán gộp.\n"
    "7.3. Xây dựng module cho quản lý bãi xe:\n"
    "- Dashboard công suất, trạng thái khu đỗ, ô EV, trụ sạc và log vào/ra.\n"
    "- Theo dõi hàng chờ, cảnh báo quá tải và các đề xuất vận hành.\n"
    "7.4. Xây dựng module cho quản trị viên:\n"
    "- Quản lý toàn mạng lưới bãi đỗ - trạm sạc.\n"
    "- Cấu hình engine dự báo, giá dịch vụ, ngưỡng guarantee và phân quyền.\n"
    "7.5. Nghiên cứu mô hình ETA Reservation & Guarantee Engine:\n"
    "- Dự báo công suất tại thời điểm người dùng đến nơi.\n"
    "- Giữ chỗ có điều kiện và tự động fallback khi tài nguyên biến động.\n"
    "7.6. Tích hợp park-and-charge cho xe điện:\n"
    "- Đặt chỗ ô EV + trụ sạc trong một quy trình thống nhất.\n"
    "- Theo dõi charging session, mức tiêu thụ điện và chi phí tổng hợp.\n"
    "7.7. Xây dựng mô hình digital twin cho vận hành:\n"
    "- Mô phỏng tái cấu hình ô đỗ thường, ô EV và vùng ưu tiên.\n"
    "- Đánh giá tác động tới công suất, doanh thu và mức độ phục vụ.\n"
    "7.8. Kiểm thử và tối ưu:\n"
    "- Kiểm thử chức năng, kiểm thử tình huống dự báo sai, quá tải, fallback và park-and-charge.\n"
    "- Tối ưu hiệu năng, trải nghiệm UI/UX và chuẩn bị dữ liệu cho giai đoạn triển khai thực tế."
)

METHOD_TEXT = (
    "Phương pháp nghiên cứu, cách tiếp cận vấn đề\n"
    "8.1. Tiếp cận vấn đề:\n"
    "- Tiếp cận hệ thống theo hướng software-first: xem bãi đỗ và trạm sạc như một mạng lưới tài nguyên có thể quan sát, dự báo và điều phối.\n"
    "- Tiếp cận bài toán từ cả hai phía: hành trình người dùng và bài toán vận hành của đơn vị khai thác.\n"
    "- Phạm vi ưu tiên vào phần mềm, dữ liệu và cơ chế ra quyết định; phần cứng thật sẽ được tích hợp ở giai đoạn sau.\n"
    "8.2. Phương pháp nghiên cứu:\n"
    "- Nghiên cứu tài liệu về smart parking, EV charging, OCPP, digital twin và ETA-based reservation.\n"
    "- Khảo sát các giải pháp công khai trong và ngoài nước để xác định khoảng trống nghiên cứu.\n"
    "- Phân tích nghiệp vụ, thiết kế use case, thiết kế cơ sở dữ liệu và kiến trúc hệ thống.\n"
    "- Phát triển prototype trên nền React + TypeScript + Node.js + SQL Server.\n"
    "- Mô phỏng dữ liệu camera, barrier, charger và kiểm thử theo kịch bản thực tế.\n"
    "- Đánh giá tính mới thông qua mô hình guarantee, fallback và digital twin trong cùng một nền tảng."
)


PROGRESS_ROWS = {
    8: [
        "Phân tích yêu cầu nghiệp vụ\nvà chốt mô hình hệ thống\nfinal",
        "Bộ tài liệu\nuse case,\nluồng nghiệp vụ,\ndanh sách chức năng lõi\nvà mô hình dữ liệu\ntổng quan",
        "01/08/2025\n-\n15/09/2025",
        "Cả nhóm",
    ],
    9: [
        "Thiết kế cơ sở dữ liệu\nvà mô hình bãi đỗ - trạm sạc",
        "Schema hoàn chỉnh:\nAppUsers,\nVehicles,\nParkingLots,\nParkingSlots,\nChargingStations,\nReservations,\nSessions,\nLogs",
        "16/09/2025\n-\n15/10/2025",
        "Cả nhóm",
    ],
    10: [
        "Thiết kế UI/UX,\ndashboard\nvà digital twin",
        "Bộ giao diện hoàn chỉnh\ncho user, manager, admin;\nmô hình số cho vận hành",
        "16/10/2025\n-\n15/12/2025",
        "Cả nhóm",
    ],
    11: [
        "Xây dựng backend,\nETA reservation engine,\npark-and-charge",
        "Prototype hoạt động được\nvới ETA booking,\nfallback, check-in,\ncharging session\nvà thanh toán gộp",
        "16/12/2025\n-\n15/05/2026",
        "Cả nhóm",
    ],
    12: [
        "Kiểm thử, tối ưu,\nhoàn thiện hồ sơ\nvà demo nghiệm thu",
        "Hệ thống hoàn chỉnh,\ntối ưu hiệu năng,\ncó kịch bản demo\nvà hồ sơ nghiên cứu\nđầy đủ",
        "16/05/2026\n-\n01/08/2026",
        "Cả nhóm",
    ],
}


TABLE3_SECTION1 = {
    2: [
        "Nền tảng Spot Ace Park - User App",
        "✅ Tìm kiếm bãi đỗ theo điểm đến, khoảng cách, giá và hỗ trợ EV\n"
        "✅ Hiển thị xác suất còn chỗ theo ETA\n"
        "✅ Đặt chỗ có điều kiện, theo dõi phiên đỗ và thanh toán gộp\n"
        "✅ Giao diện responsive, hiện đại, dễ sử dụng",
    ],
    3: [
        "Spot Ace Park - Backend & Data Platform",
        "✅ RESTful API với Node.js + Express\n"
        "✅ Cơ sở dữ liệu SQL Server với schema phục vụ bãi đỗ + trạm sạc\n"
        "✅ Quản lý booking, charging reservation, session, log và audit\n"
        "✅ Khả năng mở rộng tích hợp thiết bị thật",
    ],
    4: [
        "ETA Reservation & Guarantee Engine",
        "✅ Dự báo khả năng còn chỗ tại thời điểm người dùng đến nơi\n"
        "✅ Giữ chỗ có điều kiện theo ETA\n"
        "✅ Tự động fallback sang bãi thay thế khi bãi đích biến động\n"
        "✅ Ghi nhận và theo dõi chỉ số SLA/độ tin cậy",
    ],
    5: [
        "Manager Dashboard & Digital Twin",
        "✅ Dashboard công suất, doanh thu, trạng thái ô EV và trụ sạc\n"
        "✅ Cảnh báo quá tải và hàng chờ\n"
        "✅ Mô phỏng tái cấu hình bãi đỗ, so sánh tác động trước khi áp dụng\n"
        "✅ Hỗ trợ ra quyết định vận hành thời gian thực",
    ],
    6: [
        "Park-and-Charge Module",
        "✅ Đặt chỗ ô EV + trụ sạc trong một quy trình thống nhất\n"
        "✅ Theo dõi phiên sạc, năng lượng tiêu thụ và chi phí tổng hợp\n"
        "✅ Hỗ trợ loại đầu nối, trạng thái charger và log vận hành\n"
        "✅ Sẵn sàng mở rộng theo hướng OCPP/charger integration",
    ],
}


TABLE3_SECTION2 = {
    9: [
        "Hệ thống Spot Ace Park",
        "Hiệu năng: phản hồi nhanh cho tìm kiếm, đặt chỗ và điều phối fallback.\n"
        "Tính tiện dụng: giao diện thống nhất cho toàn hành trình đỗ xe - sạc điện.\n"
        "Tính ứng dụng: có thể triển khai dạng phần mềm trước khi tích hợp phần cứng thật.",
        "Khoa Công nghệ thông tin - Đại học Kinh doanh và Công nghệ Hà Nội.\n"
        "Báo cáo nghiệm thu đề tài, hội nghị sinh viên nghiên cứu khoa học, tạp chí khoa học của trường.",
    ],
    10: [
        "Hệ thống Park-and-Charge",
        "Độ chính xác dữ liệu: cập nhật trạng thái ô EV, trụ sạc và phiên sạc theo thời gian thực.\n"
        "Tính tương thích: hỗ trợ mô hình dữ liệu cho nhiều loại đầu nối và nhiều kịch bản sạc.\n"
        "Tính mở rộng: thuận lợi cho tích hợp trạm sạc thật trong giai đoạn sau.",
        "Báo cáo chuyên đề hoặc bài báo ứng dụng về hạ tầng sạc và dịch vụ park-and-charge.",
    ],
    11: [
        "Module Digital Twin & Simulation",
        "Mô phỏng được các kịch bản thay đổi số ô thường, ô EV, vùng ưu tiên và tải vận hành.\n"
        "Cho phép so sánh công suất, doanh thu, tỷ lệ phục vụ và mức độ quá tải trước khi áp dụng.\n"
        "Có giá trị hỗ trợ ra quyết định cho đơn vị khai thác bãi xe.",
        "Báo cáo nghiên cứu ứng dụng, chuyên đề mô phỏng vận hành đô thị thông minh.",
    ],
    12: [
        "Engine dự báo ETA, guarantee & pricing",
        "Dự báo công suất theo thời gian đến dự kiến, tính toán guarantee và cơ chế fallback.\n"
        "Kết hợp được với dynamic pricing và chỉ số SLA để tối ưu vận hành.\n"
        "Tạo nền tảng cho các nghiên cứu tiếp theo về prediction, recommendation và smart charging.",
        "Bài báo nghiên cứu ứng dụng hoặc báo cáo học thuật nội bộ về mô hình điều phối thông minh.",
    ],
}


GOAL_BLOCKS = [
    ("Normal", "Đề tài tập trung xây dựng nền tảng Spot Ace Park theo định hướng kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh trong cùng một hệ thống phần mềm thống nhất, phục vụ đồng thời nhu cầu tìm chỗ đỗ, theo dõi công suất bãi, hỗ trợ người dùng xe điện và tối ưu vận hành cho đơn vị khai thác.", WD_ALIGN_PARAGRAPH.JUSTIFY),
    ("Normal", "Mục tiêu chính của đề tài là chuyển bài toán từ mức tra cứu thông tin sang mức dự báo, điều phối và bảo đảm dịch vụ. Thay vì chỉ cho biết bãi nào đang còn chỗ tại thời điểm hiện tại, hệ thống hướng tới việc trả lời câu hỏi quan trọng hơn đối với người dùng: bãi nào có khả năng còn phù hợp khi người dùng thực sự đến nơi.", WD_ALIGN_PARAGRAPH.JUSTIFY),
    ("Normal", "Đối với người dùng (User Module):", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng ứng dụng web có giao diện trực quan, hiện đại và dễ thao tác trên cả máy tính lẫn thiết bị di động.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Cho phép tìm kiếm bãi đỗ theo điểm đến, khoảng cách, giá, loại phương tiện và mức độ hỗ trợ xe điện.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Hiển thị xác suất còn chỗ theo thời gian đến dự kiến (ETA) để người dùng có cơ sở ra quyết định chính xác hơn.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đặt chỗ có điều kiện, theo dõi thời gian giữ chỗ và nhận gợi ý bãi thay thế khi bãi đích biến động.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi phiên đỗ xe, thời gian còn lại, chi phí tạm tính và lịch sử giao dịch trên cùng một tài khoản.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Đối với người dùng xe điện:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Lọc các bãi đỗ có hỗ trợ ô EV và trụ sạc phù hợp với loại đầu nối của phương tiện.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đặt chỗ ô EV và trụ sạc trong một quy trình thống nhất theo mô hình park-and-charge.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi phiên sạc, mức điện năng tiêu thụ, công suất sạc và chi phí tổng hợp ngay trên ứng dụng.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tạo nền tảng để mở rộng sang các tính năng như queue management, idle fee và smart charging trong các giai đoạn sau.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Đối với quản lý bãi xe (Manager Module):", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi công suất bãi theo khu vực, trạng thái các ô đỗ thường, ô EV và trụ sạc theo thời gian thực.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Giám sát log xe ra/vào, trạng thái check-in/check-out và các cảnh báo bất thường từ hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Nhận gợi ý tái cấu hình bãi đỗ nhằm tăng hiệu quả phục vụ trong khung giờ cao điểm.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Có công cụ quản trị vận hành theo dạng digital twin để đánh giá các kịch bản trước khi áp dụng thật.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Đối với quản trị viên hệ thống (Admin Module):", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Quản lý toàn mạng lưới bãi đỗ - trạm sạc, tài khoản người dùng, quản lý bãi và cấu hình hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Cấu hình engine dự báo, ngưỡng guarantee, quy tắc fallback, dynamic pricing và các thông số vận hành trọng yếu.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi hiệu năng toàn hệ thống, tỷ lệ thành công của đặt chỗ, mức độ sử dụng tài nguyên và các chỉ số SLA.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Về mặt công nghệ:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Phát triển ETA Reservation & Guarantee Engine làm lõi ra quyết định của hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Áp dụng mô hình software-first để có thể mô phỏng camera, barrier, IoT logs và charger trước khi tích hợp phần cứng thật.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng kiến trúc mở cho phép mở rộng sang OCPP, smart charging, prediction engine và phân tích vận hành nâng cao.", WD_ALIGN_PARAGRAPH.LEFT),
]


OVERVIEW_BLOCKS = [
    ("Normal", "6.1. Tổng quan tình hình nghiên cứu:", None),
    ("Normal", "Trên thế giới, lĩnh vực smart parking đã phát triển theo nhiều hướng như nhận diện ô trống bằng IoT, bản đồ bãi đỗ theo thời gian thực, đặt chỗ trước, điều hướng và thanh toán điện tử. Song song với đó, lĩnh vực quản lý trạm sạc xe điện cũng phát triển mạnh với các nền tảng tìm trạm sạc, quản lý phiên sạc, cân bằng tải và kết nối thiết bị theo giao thức chuẩn.", None),
    ("Normal", "Tuy nhiên, trong nhiều trường hợp, hai bài toán đỗ xe và sạc điện vẫn tồn tại tách rời. Người dùng phải tự ghép thông tin từ nhiều nền tảng khác nhau để quyết định nơi đỗ xe, thời điểm sạc và phương án di chuyển. Điều này làm giảm tính liền mạch của trải nghiệm và khiến tài nguyên hạ tầng chưa được khai thác tối ưu.", None),
    ("Normal", "Ở Việt Nam, các giải pháp công khai hiện nay đã cho thấy nhu cầu số hóa là rất rõ. Đã xuất hiện ứng dụng bãi đỗ thời gian thực, ứng dụng bản đồ trạm sạc, các nền tảng quản lý charger và những giải pháp nội bộ cho từng bãi xe. Dù vậy, phần lớn các hệ thống mới tập trung vào một lát cắt riêng của bài toán, chưa tạo được một lớp điều phối đủ mạnh giữa người dùng và toàn bộ hạ tầng.", None),
    ("Normal", "Một điểm đáng chú ý là nhiều nền tảng mới chỉ mạnh ở việc phản ánh trạng thái hiện tại, trong khi khó khăn thực tế của người dùng lại nằm ở trạng thái tương lai gần. Người dùng quyết định dựa trên thời điểm mình sẽ đến nơi, chứ không phải chỉ dựa trên ảnh chụp hiện tại của hệ thống. Khoảng chênh giữa hai góc nhìn này chính là khoảng trống nghiên cứu mà đề tài tập trung giải quyết.", None),
    ("Normal", "Ngoài ra, phần lớn các hệ thống hiện hành chưa mô tả rõ cơ chế ứng xử khi dự báo sai. Nếu một bãi đang được khuyến nghị nhưng vài phút sau rơi vào trạng thái quá tải, hệ thống sẽ chuyển hướng thế nào, đảm bảo quyền lợi người dùng ra sao và đơn vị vận hành sẽ xử lý tình huống như thế nào. Đây là vấn đề rất quan trọng nếu muốn nâng hệ thống từ mức tham khảo lên mức dịch vụ đáng tin cậy.", None),
    ("Normal", "6.2. Tính cấp thiết của đề tài:", None),
    ("Normal", "Áp lực giao thông tại các đô thị lớn, nhu cầu đỗ xe ngày càng tăng và xu hướng phát triển xe điện đang khiến mô hình bãi đỗ truyền thống bộc lộ nhiều hạn chế. Người dùng không chỉ cần biết bãi nào còn chỗ ở hiện tại, mà cần biết bãi nào có khả năng còn phù hợp khi họ thật sự đến nơi. Đơn vị vận hành cũng không chỉ cần báo cáo cuối ngày, mà cần công cụ dự báo và điều phối ở thời gian gần như thực.", None),
    ("Normal", "Nếu không có hệ thống kết nối và dự báo, người dùng thường phải chạy vòng quanh tìm chỗ, gây lãng phí thời gian, nhiên liệu và làm tăng áp lực giao thông nội đô. Đối với xe điện, khó khăn còn lớn hơn vì ngoài chỗ đỗ còn phải cân nhắc trụ sạc, loại đầu nối, công suất và thời gian chờ.", None),
    ("Normal", "Từ phía đơn vị khai thác, việc thiếu công cụ điều phối cũng làm giảm hiệu quả sử dụng tài nguyên. Có những khung giờ bãi xe quá tải cục bộ ở một vài khu nhưng lại còn dư công suất ở khu khác; có những ngày tỷ lệ xe điện tăng mạnh nhưng số ô EV vẫn giữ như cấu hình ngày thường; có những tình huống trụ sạc hoạt động kém hiệu quả do chưa được phân bổ đúng vào nhóm người dùng phù hợp.", None),
    ("Normal", "Như vậy, tính cấp thiết của đề tài không chỉ đến từ nhu cầu của người sử dụng cuối mà còn đến từ nhu cầu chuyển đổi mô hình vận hành bãi xe. Việc hình thành một nền tảng kết nối và điều phối có thể mang lại giá trị cho nhiều bên cùng lúc, từ đó tăng khả năng ứng dụng thực tiễn của đề tài.", None),
    ("List Paragraph", "Giảm ùn tắc giao thông do giảm số lượng xe chạy vòng quanh chỉ để tìm chỗ đỗ.", None),
    ("List Paragraph", "Tiết kiệm thời gian và chi phí cho người dùng nhờ dữ liệu dự báo và đặt chỗ có điều kiện.", None),
    ("List Paragraph", "Tăng hiệu quả khai thác hạ tầng cho đơn vị vận hành thông qua điều phối bãi đỗ và tài nguyên EV.", None),
    ("List Paragraph", "Hỗ trợ xu hướng chuyển đổi số và xu hướng mở rộng hạ tầng phục vụ phương tiện sử dụng điện.", None),
    ("List Paragraph", "Tạo nền tảng dữ liệu để phát triển các nghiên cứu tiếp theo về prediction, recommendation và smart charging.", None),
    ("Normal", "6.3. Ý nghĩa và tính mới về khoa học và thực tiễn:", None),
    ("Normal", "Về mặt khoa học, đề tài chuyển trọng tâm từ ứng dụng tra cứu sang mô hình điều phối dự báo. Hệ thống không chỉ phản ánh trạng thái hiện tại mà còn dự báo công suất theo ETA, áp dụng cơ chế guarantee và fallback khi bãi đích biến động.", None),
    ("Normal", "Về mặt thực tiễn, đề tài tạo ra một nền tảng có thể phục vụ cả người dùng, xe điện, quản lý bãi xe và quản trị viên trên cùng một kiến trúc. Điểm mới nổi bật là sự kết hợp giữa reservation theo ETA, park-and-charge và digital twin trong cùng một hướng giải pháp.", None),
    ("Normal", "Về ý nghĩa triển khai, đề tài lựa chọn mô hình software-first nên có thể được thử nghiệm ở quy mô nghiên cứu mà không buộc phải đầu tư đồng thời toàn bộ phần cứng. Điều này giúp đề tài vừa phù hợp với điều kiện thực tế, vừa vẫn giữ được khả năng mở rộng sang triển khai thật trong tương lai.", None),
    ("Normal", "So với cách tiếp cận chỉ tập trung vào giao diện hoặc một tính năng đơn lẻ, hướng nghiên cứu của đề tài nhấn mạnh cơ chế ra quyết định và khả năng phối hợp nhiều dịch vụ trong cùng hệ thống. Đây là lý do làm cho đề tài có chiều sâu hơn ở cả góc độ khoa học ứng dụng lẫn góc độ phát triển phần mềm.", None),
]


RESEARCH_CONTENT_BLOCKS = [
    ("Normal", "Nội dung nghiên cứu", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.1. Phân tích và thiết kế hệ thống", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Phân tích yêu cầu chức năng và phi chức năng cho ba nhóm vai trò: người dùng, quản lý bãi xe, admin.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế use case, luồng nghiệp vụ và cơ chế phối hợp giữa tìm kiếm, đặt chỗ, check-in, phiên đỗ và phiên sạc.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế kiến trúc tổng thể theo mô hình web app + dashboard + backend điều phối + tầng dữ liệu vận hành.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế mô hình cơ sở dữ liệu cho bãi đỗ, ô đỗ, reservation, charger, charging session, IoT logs, payment và audit logs.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Mô tả rõ cơ chế tương tác giữa dữ liệu hiện trường và giao diện quản trị để làm nền cho pha triển khai phần mềm sau này.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.2. Xây dựng module cho người dùng", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Phát triển giao diện tìm kiếm bãi đỗ theo điểm đến, khoảng cách, giá và hỗ trợ EV.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng màn hình hiển thị xác suất còn chỗ theo ETA, thời gian đi bộ và điều kiện guarantee.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng quy trình đặt chỗ liền mạch: chọn bãi -> xác nhận ETA -> giữ chỗ -> check-in -> thanh toán.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng chức năng theo dõi phiên đỗ xe, nhắc thời gian và lịch sử giao dịch.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Bổ sung khả năng đánh giá bãi đỗ, xem lịch sử lựa chọn và lưu ưu tiên cá nhân của người dùng.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế cơ chế thông báo khi guarantee sắp hết hiệu lực hoặc khi cần chuyển sang phương án thay thế.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.3. Xây dựng module cho người dùng xe điện", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Lọc bãi đỗ có hỗ trợ ô EV và trụ sạc phù hợp với phương tiện.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế quy trình park-and-charge: đặt ô EV + trụ sạc + theo dõi charging session trong cùng một giao diện.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi năng lượng tiêu thụ, thời gian sạc, trạng thái charger và chi phí tổng hợp.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Ghi nhận trạng thái trụ sạc theo các mức: available, occupied, maintenance, offline để phục vụ điều phối.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đề xuất cấu trúc dữ liệu cho charging reservation, charging session và billing nhằm mở đường cho smart charging.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.4. Xây dựng module cho quản lý bãi xe", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế dashboard theo dõi công suất bãi, mật độ khu đỗ, số ô EV đang hoạt động, log xe vào/ra và cảnh báo vận hành.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Bổ sung khả năng theo dõi hàng chờ, trạng thái fallback và các đề xuất tái cấu hình bãi.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Nghiên cứu mô hình quản trị theo ca/ngày/sự kiện nhằm hỗ trợ ra quyết định vận hành.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế các chỉ số vận hành quan trọng như occupancy rate, fallback rate, guarantee success rate, EV slot utilization và charger uptime.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng cơ chế giám sát sự kiện bất thường như quá tải cục bộ, xe đến muộn, trụ sạc lỗi hoặc chênh lệch dữ liệu cảm biến.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.5. Xây dựng module cho quản trị viên", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Quản lý toàn mạng lưới bãi đỗ - trạm sạc, tài khoản và phân quyền.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Cấu hình ETA engine, guarantee rules, fallback rules, dynamic pricing và tham số hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Theo dõi hiệu năng nhiều bãi trên cùng dashboard quản trị tổng hợp.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Quản trị các tham số phục vụ nghiên cứu như cấu hình mô phỏng, ngưỡng cảnh báo, định danh khu vực và chính sách ưu tiên.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.6. Nghiên cứu ETA Reservation & Guarantee Engine", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Nghiên cứu các biến ảnh hưởng đến công suất bãi đỗ: giờ cao điểm, thời tiết, lịch sử lưu đỗ, vùng điểm đến, tỷ lệ EV.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế mô hình dự báo xác suất còn chỗ tại thời điểm người dùng đến nơi.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Xây dựng cơ chế guarantee giữ chỗ có điều kiện và fallback sang bãi thay thế khi tài nguyên biến động.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Nghiên cứu cách ghi nhận và đo lường độ tin cậy của dự báo để cải tiến engine theo thời gian.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế logic phân hạng bãi thay thế theo khoảng cách, công suất, phù hợp EV, thời gian đi bộ và khả năng phục vụ.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.7. Xây dựng mô hình digital twin cho vận hành", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Mô phỏng tái cấu hình ô đỗ thường, ô EV, vùng ưu tiên và kịch bản quá tải.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "So sánh tác động của các phương án vận hành tới công suất, doanh thu, mức độ phục vụ và khả năng khai thác.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Dùng digital twin như công cụ hỗ trợ quản lý thử nghiệm trước khi áp dụng thay đổi thực tế.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tạo cơ sở cho việc phát triển module quy hoạch đầu tư hạ tầng trong giai đoạn tiếp theo.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "7.8. Kiểm thử và tối ưu hóa", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Kiểm thử chức năng đặt chỗ, check-in, fallback, thanh toán và park-and-charge.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Kiểm thử dữ liệu mô phỏng cho camera, barrier, IoT logs và charger.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tối ưu hiệu năng tải trang, phản hồi API, tính ổn định của dashboard và khả năng mở rộng dữ liệu.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tổ chức rà soát nội dung và giao diện dựa trên phản hồi thử nghiệm để hoàn thiện hồ sơ nghiệm thu.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Khối nội dung nghiên cứu trên hướng tới việc tạo ra một hệ thống vừa có chiều sâu học thuật, vừa đủ rõ về kỹ thuật để phát triển tiếp thành phần mềm hoàn chỉnh trong giai đoạn tiếp theo.", WD_ALIGN_PARAGRAPH.LEFT),
]


METHOD_BLOCKS = [
    ("Normal", "Phương pháp nghiên cứu, cách tiếp cận vấn đề", None),
    ("Normal", "8.1. Tiếp cận vấn đề:", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Mục tiêu: Xây dựng nền tảng kết nối mạng lưới bãi đỗ xe - trạm sạc thông minh theo hướng dự báo, điều phối và bảo đảm dịch vụ.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Phạm vi: Tập trung vào phần mềm, dữ liệu và mô hình ra quyết định; phần cứng thật như barrier, camera, charger có thể tích hợp ở giai đoạn sau.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Đối tượng: Người dùng phương tiện cá nhân, người dùng xe điện, đơn vị vận hành bãi xe và quản trị viên hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Cách tiếp cận: Tiếp cận hệ thống theo hướng software-first và service orchestration. Xem bãi đỗ và trạm sạc như một mạng lưới tài nguyên có thể quan sát, dự báo, phân bổ và tái cấu hình.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "8.2. Phương pháp nghiên cứu", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Nghiên cứu tài liệu:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Khảo sát các hệ thống smart parking, EV charging, digital twin và ETA-based reservation trong và ngoài nước.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tham khảo tài liệu về OCPP, smart charging, prediction engine và quản trị hạ tầng thông minh.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Phân tích các xu hướng pháp lý, hạ tầng và nhu cầu người dùng liên quan đến bãi đỗ xe và trạm sạc.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tổng hợp các điểm mạnh, điểm hạn chế và khoảng trống nghiên cứu từ các giải pháp công khai hiện có.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Phương pháp phân tích và thiết kế:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Phân tích use case, luồng dữ liệu và tương tác giữa các vai trò trong hệ thống.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế kiến trúc hệ thống, mô hình dữ liệu và các service lõi cho đặt chỗ, fallback, charging session và dashboard.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Thiết kế cell-level content, use case và prototype để bảo đảm đề tài có thể chuyển tiếp sang thi công phần mềm.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Phương pháp phát triển sản phẩm:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Áp dụng mô hình phát triển lặp (iterative development) để cho phép tinh chỉnh dần theo phản hồi.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Ưu tiên dựng prototype chạy được trên web trước khi tính đến tích hợp phần cứng thật.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Tách rõ phần lõi nghiệp vụ và phần mô phỏng để thuận lợi khi chuyển sang môi trường triển khai thật.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Thử nghiệm và đánh giá:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Mô phỏng các kịch bản bãi đầy, guarantee thất bại, fallback sang bãi khác, queue EV và charging session.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đánh giá hệ thống theo các tiêu chí: tính mới, tính ứng dụng, khả năng mở rộng, độ liền mạch của trải nghiệm và hiệu quả vận hành.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Rà soát UI/UX, thời gian phản hồi và mức độ phù hợp với bối cảnh người dùng Việt Nam.", WD_ALIGN_PARAGRAPH.LEFT),
    ("Normal", "Phương pháp đánh giá tính mới:", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đối chiếu hệ thống đề xuất với các nhóm giải pháp hiện có để làm rõ phần đóng góp về ETA, guarantee, fallback và digital twin.", WD_ALIGN_PARAGRAPH.LEFT),
    ("List Paragraph", "Đánh giá tính khả thi trên góc nhìn phần mềm trước, sau đó mở rộng sang tích hợp hạ tầng thật ở giai đoạn tiếp theo.", WD_ALIGN_PARAGRAPH.LEFT),
]


def unique_cells(row):
    seen = []
    cells = []
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.append(key)
            cells.append(cell)
    return cells


def copy_run_style(source_run, target_run) -> None:
    if source_run is None:
        target_run.font.name = "Times New Roman"
        target_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        target_run.font.size = Pt(13)
        return
    target_run.font.name = source_run.font.name or "Times New Roman"
    target_run._element.rPr.rFonts.set(qn("w:eastAsia"), target_run.font.name)
    target_run.font.size = source_run.font.size or Pt(13)
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    if source_run.font.color is not None and source_run.font.color.rgb is not None:
        target_run.font.color.rgb = source_run.font.color.rgb


def remove_row_flags(row) -> None:
    trPr = row._tr.trPr
    if trPr is None:
        return
    for child in list(trPr):
        if child.tag in {qn("w:cantSplit"), qn("w:trHeight")}:
            trPr.remove(child)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag in {qn("w:r"), qn("w:hyperlink"), qn("w:proofErr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            p.remove(child)


def replace_cell_text(cell, new_text: str) -> None:
    paragraphs = cell.paragraphs
    base_paragraph = paragraphs[0]
    source_run = None
    for para in paragraphs:
        for run in para.runs:
            source_run = run
            break
        if source_run is not None:
            break
    for para in paragraphs[1:]:
        para._element.getparent().remove(para._element)
    clear_paragraph(base_paragraph)
    run = base_paragraph.add_run(new_text)
    copy_run_style(source_run, run)


def replace_cell_blocks(cell, blocks, font_size: float | None = None) -> None:
    paragraphs = cell.paragraphs
    template_run = None
    for para in paragraphs:
        for run in para.runs:
            template_run = run
            break
        if template_run is not None:
            break

    first_para = paragraphs[0]
    for para in paragraphs[1:]:
        para._element.getparent().remove(para._element)
    clear_paragraph(first_para)
    first_para.style = blocks[0][0]
    if blocks[0][2] is not None:
        first_para.alignment = blocks[0][2]
    first_para.paragraph_format.space_before = Pt(0)
    first_para.paragraph_format.space_after = Pt(0)
    first_para.paragraph_format.line_spacing = 1.0
    run = first_para.add_run(blocks[0][1])
    copy_run_style(template_run, run)
    if font_size is not None:
        run.font.size = Pt(font_size)

    for style_name, text, align in blocks[1:]:
        para = cell.add_paragraph()
        para.style = style_name
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(text)
        copy_run_style(template_run, run)
        if font_size is not None:
            run.font.size = Pt(font_size)


def update_document() -> Path:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(SOURCE_PATH)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SOURCE_PATH))

    # Table 2: main narrative content
    table2 = doc.tables[2]
    remove_row_flags(table2.rows[4])
    remove_row_flags(table2.rows[5])
    replace_cell_blocks(unique_cells(table2.rows[1])[0], GOAL_BLOCKS)
    replace_cell_blocks(unique_cells(table2.rows[3])[0], OVERVIEW_BLOCKS)
    replace_cell_blocks(unique_cells(table2.rows[4])[1], RESEARCH_CONTENT_BLOCKS)
    replace_cell_blocks(unique_cells(table2.rows[5])[1], METHOD_BLOCKS)

    for row_idx, values in PROGRESS_ROWS.items():
        cells = unique_cells(table2.rows[row_idx])
        replace_cell_blocks(cells[1], [("Normal", line, WD_ALIGN_PARAGRAPH.CENTER) for line in values[0].split("\n") if line.strip()], font_size=11)
        replace_cell_blocks(cells[2], [("Normal", line, WD_ALIGN_PARAGRAPH.LEFT) for line in values[1].split("\n") if line.strip()], font_size=10.5)
        replace_cell_blocks(cells[3], [("Normal", line, WD_ALIGN_PARAGRAPH.CENTER) for line in values[2].split("\n") if line.strip()], font_size=11)
        replace_cell_blocks(cells[4], [("Normal", line, WD_ALIGN_PARAGRAPH.CENTER) for line in values[3].split("\n") if line.strip()], font_size=11)

    # Table 3: expected products
    table3 = doc.tables[3]
    for row_idx, values in TABLE3_SECTION1.items():
        cells = unique_cells(table3.rows[row_idx])
        replace_cell_text(cells[1], values[0])
        blocks = [("Normal", line, WD_ALIGN_PARAGRAPH.LEFT) for line in values[1].split("\n") if line.strip()]
        replace_cell_blocks(cells[2], blocks, font_size=11)
        replace_cell_text(cells[3], "")

    for row_idx, values in TABLE3_SECTION2.items():
        cells = unique_cells(table3.rows[row_idx])
        replace_cell_text(cells[1], values[0])
        blocks_1 = [("Normal", line, WD_ALIGN_PARAGRAPH.LEFT) for line in values[1].split("\n") if line.strip()]
        blocks_2 = [("Normal", line, WD_ALIGN_PARAGRAPH.LEFT) for line in values[2].split("\n") if line.strip()]
        replace_cell_blocks(cells[2], blocks_1, font_size=11)
        replace_cell_blocks(cells[3], blocks_2, font_size=11)

    # Optional refinement for funding wording without changing numbers
    table4 = doc.tables[4]
    replace_cell_text(unique_cells(table4.rows[1])[1], "Kinh phí nhà trường:")
    replace_cell_text(unique_cells(table4.rows[2])[1], "Các nguồn vốn khác:\nKinh phí đơn vị.\nNguồn khác (hợp tác ...)")

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    print(update_document())
